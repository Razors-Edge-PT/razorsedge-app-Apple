#!/usr/bin/env python3
"""
docs/wes2_orchestrator.py

WES2 Orchestration Script
- Claude Code as implementer (plain text output)
- OpenAI/ChatGPT as strict JSON reviewer
- WES2 spec doc as source of truth
- Git as checkpoint/rollback boundary

Usage:
  python docs/wes2_orchestrator.py --spec "docs/WES2_spec.docx" --start-phase 18 --max-phases 4
"""

import argparse
import datetime
import json
import os
import re
import shlex
import subprocess
import sys
import textwrap
import time
from pathlib import Path

# ---------------------------------------------------------------------------
# Optional dependency checks
# ---------------------------------------------------------------------------
try:
    import openai
except ImportError:
    print("ERROR: openai package not installed. Run: pip install openai", file=sys.stderr)
    sys.exit(1)

try:
    import docx as _docx_module  # python-docx
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

ALLOWED_WRITE_PATTERNS = [
    re.compile(r"^lib/WES2_.*\.dart$"),
    re.compile(r"^lib/WES2_widgets/WES2_.*\.dart$"),
    re.compile(r"^test/wes2/WES2_.*\.dart$"),
    re.compile(r"^docs/wes2/"),
    re.compile(r"^\.orchestrator/wes2/"),
    re.compile(r"^docs/wes2_orchestrator\.py$"),
]

BLOCKED_WRITE_PATTERNS = [
    re.compile(r"^lib/workout_entry_screen\.dart$"),
    re.compile(r"^lib/main\.dart$"),
    re.compile(r"^lib/bb3_"),
    re.compile(r"^lib/periodization_model_utils\.dart$"),
    re.compile(r"^lib/progression_engine\.dart$"),
    re.compile(r"^pubspec\.yaml$"),
    re.compile(r"^assets/"),
    re.compile(r"^lib/generated/"),
    re.compile(r"\.g\.dart$"),
    re.compile(r"\.freezed\.dart$"),
    re.compile(r"^\.claude/settings\.local\.json$"),
    re.compile(r"^lib/templates\.dart$"),
    re.compile(r"^lib/template_model\.dart$"),
]

READ_ONLY_REFERENCE_FILES = [
    "lib/templates.dart",
    "lib/template_model.dart",
    "lib/workout_entry_screen.dart",
    "lib/bb3_week_planner.dart",
    "lib/bb3_day_panel.dart",
    "lib/bb3_models.dart",
    "lib/bb3_planned_exercise_service.dart",
    "lib/bb3_hint_service.dart",
    "lib/periodization_model_utils.dart",
    "lib/progression_engine.dart",
]

LAYOUT_RULES = textwrap.dedent("""\
    LAYOUT RULES (treat RenderFlex overflow as a blocker):
    - Avoid RenderFlex overflows.
    - Long text inside Row must use Expanded or Flexible.
    - Use maxLines and TextOverflow.ellipsis on all text widgets.
    - Avoid unconstrained Row children.
    - Do NOT use ListTile inside PopupMenuItem.
    - Test mentally at narrow 360px phone width.
""")

PHASE_MAP = {
    17: "Day Timer and Workout Summary",
    18: "Templates",
    19: "Settings Cog Dialog",
    20: "BB3 Structural Handoff",
    21: "Hint Engine Integration",
}

PHASE_GUIDANCE = {
    17: textwrap.dedent("""\
        Phase 17 — Day Timer and Workout Summary:
        - Implement the in-session day timer and end-of-session workout summary screen.
        - Timer tracks total session duration from first field entry or explicit start.
        - Summary shows exercises completed, sets logged, total volume, and session duration.
        - Summary is navigable from WES2; no destructive state changes from summary screen.
        - Write only WES2-scoped files.
        - Preserve all existing WES2 behavior.
    """),
    18: textwrap.dedent("""\
        Phase 18 — Templates:
        IMPORTANT: This is a continuation from the current WES2 repo state, not a restart.
        Write only WES2-scoped files (lib/WES2_*.dart, lib/WES2_widgets/WES2_*.dart,
        test/wes2/WES2_*.dart, docs/wes2/**, .orchestrator/wes2/**).

        REQUIRED BEHAVIORS:
        - Template loading REPLACES the current WES2 day/session rows (not appends).
        - If replacing existing rows, show a snackbar with Undo. If day is empty, no snackbar.
        - If the current WES2 day contains BB3-sourced planned rows and the user explicitly
          loads a template that replaces that day, clear ONLY the current selected BB3 planned
          day (current blockId/weekIndex/dayIndex) after verifying the WES2-visible write shape.
          Do NOT touch completed workout docs or non-current BB3 days.
          If the BB3 current-day write shape is uncertain, print BLOCKED rather than guessing.
        - After loading, blank template rows persist in wesPlannedExercises[].
          Do NOT move them to exercises[] — they are pending/planned, not completed.
        - Completed/logged values flow into exercises[] only through normal WES2 save/done pathways.
        - "Save workout to templates" stores structure only (exerciseId, name, circuitIndex,
          orderIndex, setCount). Never store actual logged values (weight, reps, rir, etc.).
        - Template picker/dialog/bottom sheet MUST be overflow-safe at 360px width.
        - Deduplicate exerciseIds: same exerciseId must not appear twice on the same date,
          even across different circuits. Silently skip duplicates on load.
        - Snapshot undo state BEFORE replacing current rows.
        - Preserve: selected date/session behavior, local draft/fast reopen behavior,
          Add Exercise/Add Circuit/Add Set/Remove Set/Delete/Replace/Move/Done behaviors.
        - Read lib/templates.dart and lib/template_model.dart for schema ONLY.
          Do NOT edit these files.
        - BB3 clearing (if needed) must only clear the current selected BB3 planned day.
          Do NOT touch completed workout docs or other BB3 days.
        - If the BB3 current-day write shape is uncertain, print BLOCKED rather than guessing.
    """),
    19: textwrap.dedent("""\
        Phase 19 — Settings Cog Dialog:
        IMPORTANT: This is a WES2-only change. Do not edit Block Planner, BB2/BB3, PMU, or
        progression engine files.

        REQUIRED BEHAVIORS:
        - Implement the exercise settings cog dialog scoped to quick in-session changes only.
        - Opened by tapping the cog beside the best-at / Avg E1RM area in WES2.
        - Fields: increments, weeklyFrequency, repPeriodizationModel, microcycle rep targets,
          RIR model, current microcycle/session RIR planned targets, progressionModel.
        - Read-only context shown in dialog: current weekly instance index, current global
          instance index (never editable, just informational).
        - activeBlockId guard: if no active block exists, show a snackbar and no-op.
          Never proceed with a save if activeBlockId is absent.
        - Save writes ONLY to:
            /planned_blocks/{uid}/blocks/{blockId}.exerciseSettings[exerciseId]
          NEVER write to plannedExerciseDetails under any circumstances.
        - Internet required for save. Do NOT queue offline. If offline, show error in dialog
          and leave dialog open so user can retry.
        - Hints do not update automatically after save. Hints update only when user triggers
          exercise refresh or global refresh. Do NOT implement hint recalculation inside
          the settings dialog save flow.
        - Validation on field unfocus and before Save. Block save while errors exist.
          - increments: positive numbers only
          - weeklyFrequency: integer 1-14
          - rep targets: parse as "reps x sets"
          - DUP Signature: parse as "min - max", numeric, min < max
        - Read-only schema inspection of BB3/Block Planner files is allowed only if needed.
          Do NOT write to those files.
        - Write only WES2-scoped files.
        - Dialog must be overflow-safe at 360px width.
    """),
    20: textwrap.dedent("""\
        Phase 20 — BB3 Structural Handoff:
        IMPORTANT: This phase writes back to the BB3 planned day Firestore path from within
        WES2 service/screen code only. Do NOT edit BB3 Dart files.

        REQUIRED BEHAVIORS:
        - Implement BB3 planned-day structural updates triggered by WES2 user actions on
          BB3-sourced rows only. WES2-only rows are NOT affected.
        - Scope is strictly: currently selected blockId / weekIndex / dayIndex.
        - Actions that must write back to BB3 planned day:
            delete a BB3-sourced exercise
            replace a BB3-sourced exercise
            move/reorder a BB3-sourced exercise between circuits or within a circuit
        - Write ONLY to the BB3 planned day path:
            /planned_blocks/{uid}/blocks/{blockId}/weeks/week_{weekIndex}/days/day_{dayIndex}
        - Verify the planned-day Firestore write shape by reading existing WES2_plan_service
          or other existing WES2 service code. Do NOT copy internal BB3 Dart write logic.
        - Do NOT edit bb3_*.dart files.
        - Do NOT touch non-current BB3 days (other weekIndex / dayIndex combinations).
        - Do NOT touch completed workout documents (/users/{uid}/workouts/{date}) except
          through existing WES2 save paths already established.
        - If the current-day BB3 write shape is uncertain from inspecting WES2 services,
          print BLOCKED rather than guessing.
        - Undo must restore both WES2 state and any queued/executed BB3 structural write.
        - Write only WES2-scoped files.
    """),
    21: textwrap.dedent("""\
        Phase 21 — Hint Engine Integration:
        HIGH RISK PHASE. Every constraint below is a hard requirement.

        REQUIRED BEHAVIORS:
        - BB3 planned overrides display as hint values ONLY. Never as actual typed values.
          Hint priority: WES2 typed actual > BB3 explicit override > model/progression hint > empty.
        - WES2 typed actual values ALWAYS win. Never overwrite a TextEditingController's text
          during hint recalculation. Caret/cursor must not jump.
        - Never overwrite actualValue or any completed set field with a hint value.
        - Use exerciseSettings (from planned_blocks/{uid}/blocks/{blockId}.exerciseSettings)
          as the authoritative source for periodization and progression model configuration.
        - Hints recalculate instantly and locally when any field changes. No save is required
          to trigger recalculation. Recalculation must not lag or show a spinner.
        - Set 2+ hints derive from Set 1 target/actual values plus fatigue/dropoff logic.
          They do NOT independently re-run the full periodization model.
        - Do NOT store hint snapshots as completed Firestore actuals under any path.
          Hint snapshots may only exist in local Isar storage for fast reopen.
        - Read existing WES, PMU (periodization_model_utils.dart), and progression engine
          (progression_engine.dart) files for reference and schema understanding ONLY.
          Do NOT edit those files.
        - If PMU or progression engine integration requires writing to a blocked file,
          print BLOCKED rather than editing the file.
        - Do NOT auto-commit this phase unless --auto-commit-hints is explicitly passed.
          Leave changes uncommitted for human manual testing and hint verification.
        - If RIR hint changes from plan due to locked weight/reps, show a subtle visual cue
          behind/around the RIR field. Do NOT rewrite the RIR controller text.
        - Write only WES2-scoped files.
    """),
}

COMPACT_SPEC_BRIEF = textwrap.dedent("""\
    WES2 SPECIFICATION BRIEF (authoritative — see full spec for detail):

    ROLES:
    - WES2 = execution/logging surface. Athletes enter actual values here.
    - BB3 = planning surface. BB3 stores planned set hints/overrides and structure.
    - Planned values = intent. Show as hint text in WES2 only. NEVER populate controllers.
    - Completed values = reality. User-entered values in WES2 are authoritative.

    CORE RULES:
    - BB3 planned values are hints only. Priority: WES2 typed > BB3 override > model hint > empty.
    - Duplicate exerciseIds on the same date are NEVER allowed.
    - Blank WES2-added rows persist in wesPlannedExercises[].
    - Completed rows live in exercises[] only.
    - isMarkedDone lives in exercises[] only.
    - ExerciseSettings writes only to planned_blocks/{uid}/blocks/{blockId}.exerciseSettings.
    - NEVER write to plannedExerciseDetails.
    - ExerciseSettings save requires internet; do NOT queue offline.
    - BB3 structural updates (delete/replace/reorder) write back to BB3 planned day path only.
    - WES2 must use Isar for local draft/fast reopen/offline.
    - Every path must use UserContext.currentUid / selected athlete UID, not raw FirebaseAuth.

    WRITE SCOPE — READ THIS CAREFULLY BEFORE REVIEWING:

    ALLOWED WRITES — the rule is: any file whose name starts with WES2_ is allowed.
    The pattern lib/WES2_*.dart means ANY file in lib/ whose filename begins with WES2_.
    The pattern lib/WES2_widgets/WES2_*.dart means ANY file in lib/WES2_widgets/ whose
    filename begins with WES2_.

    Currently known WES2 files (ALL of these are fully allowed to be read AND written):
      lib/WES2_screen.dart
      lib/WES2_controller.dart
      lib/WES2_models.dart
      lib/WES2_repository.dart          <-- ALLOWED. Do NOT reject edits to this file.
      lib/WES2_plan_service.dart
      lib/WES2_local_store.dart
      lib/WES2_template_service.dart
      lib/WES2_sync_service.dart
      lib/WES2_hint_service.dart
      lib/WES2_widgets/WES2_day_actions_row.dart
      lib/WES2_widgets/WES2_day_header.dart
      lib/WES2_widgets/WES2_empty_state.dart
      lib/WES2_widgets/WES2_exercise_card.dart
      lib/WES2_widgets/WES2_exercise_picker.dart
      lib/WES2_widgets/WES2_field_cell.dart
      lib/WES2_widgets/WES2_set_row.dart
      test/wes2/WES2_*.dart  (any test file starting with WES2_)
      docs/wes2/**
      .orchestrator/wes2/**
      docs/wes2_orchestrator.py

    New files are also allowed if their filename starts with WES2_ and they live in
    lib/, lib/WES2_widgets/, or test/wes2/.

    BLOCKED WRITES — reject ONLY if the plan writes one of these:
      lib/workout_entry_screen.dart
      lib/main.dart
      lib/bb3_*.dart  (any file starting with bb3_)
      lib/periodization_model_utils.dart
      lib/progression_engine.dart
      pubspec.yaml
      assets/**
      *.g.dart  (generated)
      *.freezed.dart  (generated)
      .claude/settings.local.json
      lib/templates.dart
      lib/template_model.dart

    Read-only reference (Claude may read but NOT edit):
      lib/templates.dart, lib/template_model.dart, lib/workout_entry_screen.dart,
      lib/bb3_*.dart, lib/periodization_model_utils.dart, lib/progression_engine.dart

    CRITICAL REVIEWER REMINDER:
    lib/WES2_repository.dart IS in the allowed write scope. Do NOT reject it.
    Any lib/WES2_*.dart file is allowed. The WES2_ prefix is the only test.
    Do NOT reject a plan for reading a reference file — only for writing a blocked file.
""")


# ---------------------------------------------------------------------------
# Helper: read spec docx
# ---------------------------------------------------------------------------

def read_spec_docx(spec_path: str) -> str:
    """Read full text from a .docx spec file."""
    import docx as _docx
    doc = _docx.Document(spec_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also include table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in paragraphs:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def extract_phase_excerpt(full_spec: str, phase_no: int, phase_name: str, max_chars: int = 4000) -> str:
    """Extract spec text relevant to a phase using keyword matching."""
    phase_name_lower = phase_name.lower()
    keywords = set(phase_name_lower.split())

    # Add phase-specific keywords
    extra = {
        17: {"timer", "summary", "duration", "session"},
        18: {"template", "templates", "loading", "load", "replace", "save workout"},
        19: {"settings", "cog", "exercisesettings", "increment", "frequency", "rir", "progression"},
        20: {"bb3", "handoff", "structural", "planned", "delete", "replace", "reorder"},
        21: {"hint", "hints", "periodization", "progression", "e1rm", "model", "recalcul"},
    }
    keywords |= extra.get(phase_no, set())

    lines = full_spec.split("\n")
    scored = []
    for i, line in enumerate(lines):
        line_lower = line.lower()
        score = sum(1 for kw in keywords if kw in line_lower)
        if score > 0:
            # Include surrounding context lines
            start = max(0, i - 2)
            end = min(len(lines), i + 5)
            scored.append((score, i, "\n".join(lines[start:end])))

    # Deduplicate overlapping chunks and sort by score
    scored.sort(key=lambda x: -x[0])
    seen_lines = set()
    excerpt_parts = []
    total_chars = 0
    for _score, line_idx, chunk in scored:
        if line_idx in seen_lines:
            continue
        seen_lines.add(line_idx)
        if total_chars + len(chunk) > max_chars:
            break
        excerpt_parts.append(chunk)
        total_chars += len(chunk)

    if not excerpt_parts:
        # Fallback: first 2000 chars of spec
        return full_spec[:2000]
    return "\n---\n".join(excerpt_parts)


# ---------------------------------------------------------------------------
# Helper: git operations
# ---------------------------------------------------------------------------

def get_repo_root() -> str:
    result = subprocess.run(
        ["git", "rev-parse", "--show-toplevel"],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print("ERROR: Not inside a git repository.", file=sys.stderr)
        sys.exit(1)
    return result.stdout.strip()


def require_clean_tree():
    result = subprocess.run(
        ["git", "status", "--porcelain"],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print("ERROR: git status failed.", file=sys.stderr)
        sys.exit(1)
    if result.stdout.strip():
        print("ERROR: Git working tree is not clean. Please commit or stash changes before running the orchestrator.")
        print("Dirty files:")
        print(result.stdout)
        sys.exit(1)


def get_changed_files() -> list:
    """Return list of files changed since last commit (unstaged + staged + untracked)."""
    result = subprocess.run(
        ["git", "status", "--porcelain"],
        capture_output=True, text=True
    )
    files = []
    for line in result.stdout.splitlines():
        if len(line) >= 3:
            filepath = line[3:].strip()
            # Handle renamed files (format: "old -> new")
            if " -> " in filepath:
                filepath = filepath.split(" -> ")[-1].strip()
            files.append(filepath)
    return files


def get_diff_text(files: list) -> str:
    """Get git diff for specific files."""
    if not files:
        return ""
    result = subprocess.run(
        ["git", "diff", "HEAD", "--"] + files,
        capture_output=True, text=True
    )
    staged = subprocess.run(
        ["git", "diff", "--cached", "--"] + files,
        capture_output=True, text=True
    )
    diff = result.stdout + staged.stdout
    # Also include untracked new files
    for f in files:
        if os.path.exists(f):
            untracked = subprocess.run(
                ["git", "ls-files", "--error-unmatch", f],
                capture_output=True, text=True
            )
            if untracked.returncode != 0:  # untracked
                try:
                    with open(f, "r", encoding="utf-8", errors="replace") as fh:
                        content = fh.read()
                    diff += f"\n--- /dev/null\n+++ b/{f}\n{content}\n"
                except Exception:
                    pass
    return diff


def git_revert_file(filepath: str):
    """Revert a tracked file or remove an untracked file."""
    result = subprocess.run(
        ["git", "ls-files", "--error-unmatch", filepath],
        capture_output=True, text=True
    )
    if result.returncode == 0:
        subprocess.run(["git", "checkout", "HEAD", "--", filepath], capture_output=True)
    else:
        try:
            os.remove(filepath)
        except FileNotFoundError:
            pass


def git_commit(files: list, message: str, body: str = "") -> str:
    """Stage and commit specific files. Returns commit hash."""
    if not files:
        return ""
    subprocess.run(["git", "add", "--"] + files, check=True)
    full_message = message
    if body:
        full_message = message + "\n\n" + body
    result = subprocess.run(
        ["git", "commit", "-m", full_message],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        return ""
    # Get commit hash
    hash_result = subprocess.run(
        ["git", "rev-parse", "--short", "HEAD"],
        capture_output=True, text=True
    )
    return hash_result.stdout.strip()


# ---------------------------------------------------------------------------
# Helper: scope enforcement
# ---------------------------------------------------------------------------

def is_allowed_write(filepath: str) -> bool:
    for pattern in ALLOWED_WRITE_PATTERNS:
        if pattern.match(filepath):
            return True
    return False


def is_blocked_write(filepath: str) -> bool:
    for pattern in BLOCKED_WRITE_PATTERNS:
        if pattern.search(filepath):
            return True
    return False


def is_orchestrator_log_file(filepath: str) -> bool:
    """
    Returns True for files under .orchestrator/wes2/.
    These are allowed to exist (scope-wise) but must be excluded from:
    - the implementation diff sent to the reviewer
    - the Dart analyzer file list
    - the auto-commit file list
    """
    return filepath.startswith(".orchestrator/wes2/")


def enforce_write_scope(changed_files: list, revert: bool, log_func) -> tuple:
    """
    Check changed files against write scope.
    Returns (allowed_files, out_of_scope_files, blocked_files).
    Reverts/removes out-of-scope files if revert=True.
    """
    allowed = []
    out_of_scope = []
    blocked = []

    for f in changed_files:
        if is_blocked_write(f):
            blocked.append(f)
            if revert:
                log_func(f"  REVERTING blocked file: {f}")
                git_revert_file(f)
        elif is_allowed_write(f):
            allowed.append(f)
        else:
            out_of_scope.append(f)
            if revert:
                log_func(f"  REVERTING out-of-scope file: {f}")
                git_revert_file(f)

    return allowed, out_of_scope, blocked


# ---------------------------------------------------------------------------
# Helper: Claude invocation
# ---------------------------------------------------------------------------

def run_claude(prompt: str, claude_cmd: str, log_dir: Path, label: str, timeout_seconds: int = 2400) -> str:
    """Run Claude with a prompt via stdin pipe. Returns plain text output."""
    log_path = log_dir / f"{label}_prompt.txt"
    log_path.write_text(prompt, encoding="utf-8")

    # Use shlex.split to correctly handle quoted args in the command string
    cmd_parts = shlex.split(claude_cmd)

    try:
        result = subprocess.run(
            cmd_parts,
            input=prompt,
            capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=timeout_seconds,
        )
        output = result.stdout or ""
        if result.stderr:
            output += "\n[STDERR]\n" + result.stderr
    except subprocess.TimeoutExpired:
        output = f"ERROR: Claude timed out after {timeout_seconds} seconds."
    except FileNotFoundError:
        output = f"ERROR: Claude command not found: {claude_cmd}"

    out_path = log_dir / f"{label}_output.txt"
    out_path.write_text(output, encoding="utf-8")
    return output


# ---------------------------------------------------------------------------
# Helper: OpenAI reviewer
# ---------------------------------------------------------------------------

def call_openai_reviewer(
    prompt: str,
    openai_model: str,
    log_dir: Path,
    label: str,
    max_retries: int = 2,
) -> dict:
    """
    Call OpenAI reviewer using the Responses API (client.responses.create).
    Returns parsed JSON dict.
    If JSON is invalid, asks once to repair using a follow-up Responses call.
    Raises RuntimeError on unrecoverable failure.

    Uses Responses API because:
    - client.responses.create supports gpt-5.4-mini and newer models natively.
    - resp.output_text gives a clean single string, no choices[0] indexing needed.
    - The system instruction is passed as the 'instructions' parameter.
    """
    log_path = log_dir / f"{label}_prompt.txt"
    log_path.write_text(prompt, encoding="utf-8")

    client = openai.OpenAI(api_key=os.environ["OPENAI_API_KEY"])

    system_instructions = textwrap.dedent("""\
        You are a strict WES2 specification guardian reviewer for a Flutter/Firebase app.
        You review implementation plans and code diffs against the WES2 specification.

        You MUST respond with ONLY valid JSON. No markdown, no backticks, no preamble.
        No explanation outside the JSON object.

        JSON schema:
        {
          "approved": true or false,
          "reason": "brief explanation",
          "spec_references_used": ["section or rule cited from spec"],
          "corrections": ["optional list of required corrections if rejected"]
        }

        ====================================================================
        WRITE SCOPE — YOU MUST APPLY THIS RULE BEFORE REJECTING ANYTHING:
        ====================================================================

        ALLOWED WRITES — the rule is simple: any file whose filename starts with WES2_ is allowed.

        These files are ALL fully allowed to be written. Do NOT reject any of them:
          lib/WES2_screen.dart
          lib/WES2_controller.dart
          lib/WES2_models.dart
          lib/WES2_repository.dart       <-- ALLOWED. Never reject edits to this file.
          lib/WES2_plan_service.dart
          lib/WES2_local_store.dart
          lib/WES2_template_service.dart
          lib/WES2_sync_service.dart
          lib/WES2_hint_service.dart
          lib/WES2_widgets/WES2_day_actions_row.dart
          lib/WES2_widgets/WES2_day_header.dart
          lib/WES2_widgets/WES2_empty_state.dart
          lib/WES2_widgets/WES2_exercise_card.dart
          lib/WES2_widgets/WES2_exercise_picker.dart
          lib/WES2_widgets/WES2_field_cell.dart
          lib/WES2_widgets/WES2_set_row.dart

        Any NEW file is also allowed if its name starts with WES2_ and it lives in
        lib/, lib/WES2_widgets/, or test/wes2/.

        BLOCKED WRITES — only reject for these:
          lib/workout_entry_screen.dart, lib/main.dart, lib/bb3_*.dart,
          lib/periodization_model_utils.dart, lib/progression_engine.dart,
          pubspec.yaml, assets/**, *.g.dart, *.freezed.dart,
          .claude/settings.local.json, lib/templates.dart, lib/template_model.dart

        READING reference files (lib/templates.dart, lib/bb3_*.dart, etc.) is ALLOWED.
        Only reject for WRITING a blocked file, never for reading one.

        ====================================================================
        BEFORE YOU REJECT — ask yourself:
        1. Does the filename start with WES2_? If yes → it is ALLOWED. Do not reject.
        2. Is it in the blocked list above? If no → it is not your reason to reject.
        ====================================================================

        spec_references_used must be non-empty whenever you reject or require corrections.
        Approve if the plan/implementation follows WES2 spec and stays within allowed scope.
    """)

    last_raw = ""
    for attempt in range(max_retries):
        try:
            response = client.responses.create(
                model=openai_model,
                instructions=system_instructions,
                input=prompt,
            )
            raw = response.output_text.strip()
            last_raw = raw
        except Exception as e:
            raise RuntimeError(f"OpenAI Responses API error: {e}")

        # Try to parse JSON
        parsed = _try_parse_json(raw)
        if parsed is not None:
            out_path = log_dir / f"{label}_response.json"
            out_path.write_text(json.dumps(parsed, indent=2), encoding="utf-8")
            return parsed

        # On first failure, ask reviewer to repair JSON
        if attempt == 0:
            repair_input = (
                "Your previous response was not valid JSON. "
                "Respond ONLY with the corrected JSON object, no other text:\n\n"
                + raw
            )
            try:
                repair_response = client.responses.create(
                    model=openai_model,
                    instructions=system_instructions,
                    input=repair_input,
                )
                raw = repair_response.output_text.strip()
                last_raw = raw
                parsed = _try_parse_json(raw)
                if parsed is not None:
                    out_path = log_dir / f"{label}_response.json"
                    out_path.write_text(json.dumps(parsed, indent=2), encoding="utf-8")
                    return parsed
            except Exception as e:
                raise RuntimeError(f"OpenAI repair attempt failed: {e}")

    err_path = log_dir / f"{label}_invalid_response.txt"
    err_path.write_text(last_raw, encoding="utf-8")
    raise RuntimeError(
        f"OpenAI reviewer returned invalid JSON after {max_retries} attempts. "
        f"Raw response saved to: {err_path}"
    )


def _try_parse_json(text: str):
    """Try to parse JSON, stripping markdown fences if present."""
    text = text.strip()
    # Strip markdown code fences
    if text.startswith("```"):
        lines = text.splitlines()
        lines = [l for l in lines if not l.startswith("```")]
        text = "\n".join(lines).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return None


def require_spec_refs_review(
    review: dict,
    openai_model: str,
    original_prompt: str,
    log_dir: Path,
    label: str,
) -> dict:
    """
    If reviewer rejected without spec_references_used, ask once to add them.
    If still empty, return review as-is with a logged warning.
    """
    if review.get("approved", True):
        return review  # Only enforce on rejections

    refs = review.get("spec_references_used", [])
    if refs:
        return review

    # Ask once
    repair_prompt = (
        "Your review JSON is missing spec_references_used. "
        "Please revise your review to include specific spec section references. "
        "Return ONLY the corrected JSON object:\n\n"
        + json.dumps(review)
    )
    try:
        repaired = call_openai_reviewer(
            original_prompt + "\n\nPREVIOUS REVIEW (missing spec refs):\n" + json.dumps(review) + "\n\n" + repair_prompt,
            openai_model=openai_model,
            log_dir=log_dir,
            label=label + "_spec_ref_repair",
        )
        if repaired.get("spec_references_used"):
            return repaired
    except Exception:
        pass

    # Log warning and return original
    warn_path = log_dir / f"{label}_missing_spec_refs.txt"
    warn_path.write_text(
        "WARNING: Reviewer failed to provide spec_references_used.\n"
        + json.dumps(review, indent=2),
        encoding="utf-8",
    )
    print("  WARNING: Reviewer rejected without providing spec references. Treating as rejection.")
    return review


# ---------------------------------------------------------------------------
# Helper: Dart analyzer
# ---------------------------------------------------------------------------

def run_dart_analyzer(dart_files: list, log_dir: Path) -> tuple:
    """
    Run dart format and flutter analyze on changed Dart files.
    Returns (success: bool, output: str).
    """
    if not dart_files:
        return True, "No Dart files to analyze."

    output_lines = []

    # dart format
    fmt_result = subprocess.run(
        ["dart", "format"] + dart_files,
        capture_output=True, text=True
    )
    output_lines.append("=== dart format ===")
    output_lines.append(fmt_result.stdout)
    if fmt_result.stderr:
        output_lines.append(fmt_result.stderr)

    # flutter analyze (targeted)
    analyze_result = subprocess.run(
        ["flutter", "analyze"] + dart_files,
        capture_output=True, text=True
    )
    output_lines.append("=== flutter analyze ===")
    output_lines.append(analyze_result.stdout)
    if analyze_result.stderr:
        output_lines.append(analyze_result.stderr)

    full_output = "\n".join(output_lines)
    log_path = log_dir / "analyzer_output.txt"
    log_path.write_text(full_output, encoding="utf-8")

    success = analyze_result.returncode == 0
    return success, full_output


# ---------------------------------------------------------------------------
# Prompt builders
# ---------------------------------------------------------------------------

def build_plan_prompt(
    phase_no: int,
    phase_name: str,
    spec_excerpt: str,
    previous_feedback: str = "",
) -> str:
    guidance = PHASE_GUIDANCE.get(phase_no, "No specific guidance available. Follow WES2 spec.")
    feedback_section = ""
    if previous_feedback:
        feedback_section = f"\n\nPREVIOUS REVIEWER FEEDBACK (you MUST address all corrections):\n{previous_feedback}\n"

    return textwrap.dedent(f"""\
        You are Claude Code implementing Phase {phase_no} ({phase_name}) of the WES2 Flutter app.

        {COMPACT_SPEC_BRIEF}

        PHASE-SPECIFIC GUIDANCE:
        {guidance}

        {LAYOUT_RULES}

        RELEVANT SPEC EXCERPT:
        {spec_excerpt}
        {feedback_section}

        TASK:
        Write a detailed implementation plan for Phase {phase_no} ({phase_name}).

        The plan must:
        1. List each WES2 file you will CREATE or MODIFY (must all be in allowed write scope).
        2. Describe exactly what each file will contain or change.
        3. Confirm which reference files you will READ (not edit) for schema context.
        4. Confirm you will NOT edit any blocked files.
        5. Address the phase-specific guidance requirements point by point.
        6. Note any BLOCKED conditions where human intervention is needed.

        If you are uncertain about a BB3 write shape or schema that would require editing a blocked file,
        write BLOCKED and explain why, rather than guessing.

        Output the plan as plain text. Be specific and thorough.
    """)


def build_plan_review_prompt(
    phase_no: int,
    phase_name: str,
    spec_excerpt: str,
    plan_text: str,
) -> str:
    guidance = PHASE_GUIDANCE.get(phase_no, "")
    return textwrap.dedent(f"""\
        You are reviewing an implementation plan for Phase {phase_no} ({phase_name}) of the WES2 Flutter app.

        {COMPACT_SPEC_BRIEF}

        PHASE-SPECIFIC REQUIREMENTS:
        {guidance}

        {LAYOUT_RULES}

        RELEVANT SPEC EXCERPT:
        {spec_excerpt}

        IMPLEMENTATION PLAN TO REVIEW:
        {plan_text}

        Review the plan against the spec and write scope rules.
        Approve if the plan correctly addresses all phase requirements and stays in scope.
        Reject if it misses critical requirements, proposes writing blocked files, or violates spec.

        Remember: Reading reference files (lib/templates.dart, lib/bb3_*.dart, etc.) is ALLOWED.
        Only reject for WRITING blocked files or violating spec behavior rules.

        Return ONLY valid JSON with no other text:
        {{
          "approved": true or false,
          "reason": "explanation",
          "spec_references_used": ["list of spec sections/rules cited"],
          "corrections": ["required corrections if rejected"]
        }}
    """)


def build_implementation_prompt(
    phase_no: int,
    phase_name: str,
    spec_excerpt: str,
    approved_plan: str,
    previous_feedback: str = "",
    out_of_scope_files: list = None,
    blocked_files: list = None,
) -> str:
    guidance = PHASE_GUIDANCE.get(phase_no, "")
    feedback_section = ""
    if previous_feedback:
        feedback_section = f"\n\nPREVIOUS REVIEWER FEEDBACK (you MUST address all corrections):\n{previous_feedback}\n"

    scope_warning = ""
    if out_of_scope_files or blocked_files:
        scope_warning = "\n\nWARNING FROM PREVIOUS ATTEMPT:\n"
        if blocked_files:
            scope_warning += f"  BLOCKED files written (REVERTED): {blocked_files}\n"
        if out_of_scope_files:
            scope_warning += f"  Out-of-scope files written (REVERTED): {out_of_scope_files}\n"
        scope_warning += "  Do NOT write to these files.\n"

    return textwrap.dedent(f"""\
        You are Claude Code implementing Phase {phase_no} ({phase_name}) of the WES2 Flutter app.

        {COMPACT_SPEC_BRIEF}

        PHASE-SPECIFIC GUIDANCE:
        {guidance}

        {LAYOUT_RULES}

        RELEVANT SPEC EXCERPT:
        {spec_excerpt}

        APPROVED IMPLEMENTATION PLAN:
        {approved_plan}
        {feedback_section}{scope_warning}

        TASK:
        Implement Phase {phase_no} ({phase_name}) now.

        Rules:
        - Write ONLY files in the allowed scope (lib/WES2_*.dart, lib/WES2_widgets/WES2_*.dart,
          test/wes2/WES2_*.dart, docs/wes2/**, .orchestrator/wes2/**).
        - Any new Dart implementation file MUST include WES2 in the filename.
        - Do NOT edit blocked files under any circumstances.
        - If you need schema from a reference file, you may read it but MUST NOT edit it.
        - If you encounter a situation requiring edits to blocked files, print BLOCKED and stop.
        - Follow the approved plan exactly.
        - Follow all layout rules (overflow-safe at 360px width).

        Implement now. Write the actual Dart code to the files.
        If BLOCKED, write "BLOCKED: <explanation>" as your entire response.
    """)


def build_implementation_review_prompt(
    phase_no: int,
    phase_name: str,
    spec_excerpt: str,
    diff_text: str,
    analyzer_output: str,
    out_of_scope_files: list,
    blocked_files: list,
    previous_feedback: str = "",
) -> str:
    guidance = PHASE_GUIDANCE.get(phase_no, "")
    scope_section = ""
    if blocked_files:
        scope_section += f"\nBLOCKED FILES WRITTEN (already reverted): {blocked_files}\n"
    if out_of_scope_files:
        scope_section += f"OUT-OF-SCOPE FILES WRITTEN (already reverted): {out_of_scope_files}\n"

    feedback_section = ""
    if previous_feedback:
        feedback_section = f"\nPREVIOUS FEEDBACK THAT MUST BE ADDRESSED:\n{previous_feedback}\n"

    # Truncate diff to avoid token limits
    diff_excerpt = diff_text[:6000] if len(diff_text) > 6000 else diff_text

    return textwrap.dedent(f"""\
        You are reviewing an implementation diff for Phase {phase_no} ({phase_name}) of the WES2 Flutter app.

        {COMPACT_SPEC_BRIEF}

        PHASE-SPECIFIC REQUIREMENTS:
        {guidance}

        {LAYOUT_RULES}

        RELEVANT SPEC EXCERPT:
        {spec_excerpt}
        {scope_section}{feedback_section}

        IMPLEMENTATION DIFF:
        {diff_excerpt}

        ANALYZER OUTPUT:
        {analyzer_output}

        Review the implementation against the spec and scope rules.
        Approve if the implementation correctly satisfies all phase requirements.
        Reject if it violates spec behavior, writes blocked files, has analyzer errors,
        has RenderFlex overflow risks, or misses critical requirements.

        Return ONLY valid JSON with no other text:
        {{
          "approved": true or false,
          "reason": "explanation",
          "spec_references_used": ["list of spec sections/rules cited"],
          "corrections": ["required corrections if rejected"]
        }}
    """)


# ---------------------------------------------------------------------------
# Logging setup
# ---------------------------------------------------------------------------

def setup_log_dir(base: str = ".orchestrator/wes2") -> Path:
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    log_dir = Path(base) / ts
    log_dir.mkdir(parents=True, exist_ok=True)
    return log_dir


class Logger:
    def __init__(self, log_dir: Path):
        self.log_dir = log_dir
        self.log_file = log_dir / "orchestrator.log"
        self._fh = open(self.log_file, "w", encoding="utf-8")

    def log(self, msg: str):
        ts = datetime.datetime.now().strftime("%H:%M:%S")
        line = f"[{ts}] {msg}"
        print(line)
        self._fh.write(line + "\n")
        self._fh.flush()

    def close(self):
        self._fh.close()


# ---------------------------------------------------------------------------
# Main orchestration loop
# ---------------------------------------------------------------------------

def run_phase(
    phase_no: int,
    full_spec: str,
    args,
    log_dir: Path,
    logger: Logger,
    start_time: float = 0.0,
) -> bool:
    """
    Orchestrate a single phase: plan -> review -> implement -> review -> commit.
    Returns True if phase completed successfully, False otherwise.
    start_time: time.monotonic() value from main(); used for total runtime guard.
    """

    def _check_runtime() -> bool:
        """Returns True if we are still within max runtime."""
        if start_time <= 0:
            return True
        elapsed = time.monotonic() - start_time
        if elapsed >= args.max_runtime_seconds:
            logger.log(
                f"\n  TIMEOUT: Total runtime {elapsed:.0f}s exceeded "
                f"--max-runtime-seconds {args.max_runtime_seconds}. Stopping safely."
            )
            return False
        return True

    phase_name = PHASE_MAP.get(phase_no, f"Phase {phase_no}")
    logger.log(f"\n{'='*60}")
    logger.log(f"PHASE {phase_no}: {phase_name}")
    logger.log(f"{'='*60}")

    spec_excerpt = extract_phase_excerpt(full_spec, phase_no, phase_name)
    excerpt_path = log_dir / f"phase_{phase_no}_spec_excerpt.txt"
    excerpt_path.write_text(spec_excerpt, encoding="utf-8")

    # ----------------------------------------------------------------
    # STEP 1-3: Plan loop
    # ----------------------------------------------------------------
    approved_plan = None
    previous_plan_feedback = ""

    for plan_attempt in range(1, args.max_plan_attempts + 1):
        if not _check_runtime():
            return False

        logger.log(f"\n  [Plan attempt {plan_attempt}/{args.max_plan_attempts}]")

        plan_prompt = build_plan_prompt(
            phase_no, phase_name, spec_excerpt,
            previous_feedback=previous_plan_feedback,
        )

        plan_label = f"phase_{phase_no}_plan_attempt_{plan_attempt}"
        logger.log("  Calling Claude for implementation plan...")
        plan_text = run_claude(
            plan_prompt, args.claude_cmd, log_dir, plan_label,
            timeout_seconds=args.claude_timeout_seconds,
        )

        # Check for BLOCKED
        if _is_blocked(plan_text):
            logger.log(f"\n  BLOCKED: human intervention required.")
            logger.log(f"  Claude output: {plan_text[:500]}")
            logger.log(f"  Log directory: {log_dir}")
            return False

        if not _check_runtime():
            return False

        logger.log("  Calling OpenAI to review plan...")
        review_prompt = build_plan_review_prompt(
            phase_no, phase_name, spec_excerpt, plan_text
        )
        review_label = f"phase_{phase_no}_plan_review_{plan_attempt}"

        try:
            review = call_openai_reviewer(
                review_prompt, args.openai_model, log_dir, review_label
            )
        except RuntimeError as e:
            logger.log(f"  ERROR: OpenAI reviewer failed: {e}")
            return False

        # Require spec refs on rejection
        if not review.get("approved", False):
            review = require_spec_refs_review(
                review, args.openai_model, review_prompt, log_dir, review_label
            )

        logger.log(f"  Plan review result: approved={review.get('approved')}")
        logger.log(f"  Reason: {review.get('reason', '')}")

        if review.get("approved", False):
            approved_plan = plan_text
            logger.log("  Plan APPROVED.")
            break
        else:
            corrections = review.get("corrections", [])
            spec_refs = review.get("spec_references_used", [])
            logger.log(f"  Plan REJECTED. Corrections: {corrections}")
            logger.log(f"  Spec refs: {spec_refs}")
            previous_plan_feedback = (
                f"Review JSON:\n{json.dumps(review, indent=2)}\n\n"
                f"You MUST address all corrections before resubmitting."
            )

    if approved_plan is None:
        logger.log(f"\n  FAILED: Plan not approved after {args.max_plan_attempts} attempts.")
        return False

    # ----------------------------------------------------------------
    # STEP 4-8: Implementation loop
    # ----------------------------------------------------------------
    previous_impl_feedback = ""
    out_of_scope_from_prev = []
    blocked_from_prev = []

    for impl_attempt in range(1, args.max_implementation_attempts + 1):
        if not _check_runtime():
            return False

        logger.log(f"\n  [Implementation attempt {impl_attempt}/{args.max_implementation_attempts}]")

        impl_prompt = build_implementation_prompt(
            phase_no, phase_name, spec_excerpt, approved_plan,
            previous_feedback=previous_impl_feedback,
            out_of_scope_files=out_of_scope_from_prev,
            blocked_files=blocked_from_prev,
        )

        impl_label = f"phase_{phase_no}_impl_attempt_{impl_attempt}"
        logger.log("  Calling Claude for implementation...")
        impl_output = run_claude(
            impl_prompt, args.claude_cmd, log_dir, impl_label,
            timeout_seconds=args.claude_timeout_seconds,
        )

        # Check for BLOCKED
        if _is_blocked(impl_output):
            logger.log(f"\n  BLOCKED: human intervention required.")
            logger.log(f"  Claude output: {impl_output[:500]}")
            logger.log(f"  Log directory: {log_dir}")
            # Scope enforcement before exit
            changed = get_changed_files()
            if changed and args.revert_out_of_scope:
                logger.log("  Running scope enforcement before BLOCKED exit...")
                enforce_write_scope(changed, revert=True, log_func=logger.log)
            return False

        # ---- Scope enforcement ----
        changed_files = get_changed_files()
        logger.log(f"  Changed files detected: {changed_files}")

        allowed_files, out_of_scope_files, blocked_files = enforce_write_scope(
            changed_files,
            revert=args.revert_out_of_scope,
            log_func=logger.log,
        )

        out_of_scope_from_prev = out_of_scope_files
        blocked_from_prev = blocked_files

        if blocked_files:
            logger.log(f"  SCOPE VIOLATION — blocked files written: {blocked_files}")
        if out_of_scope_files:
            logger.log(f"  SCOPE VIOLATION — out-of-scope files written: {out_of_scope_files}")

        # ---- Exclude orchestrator log files from code review surfaces ----
        # Log files may appear in git status if Claude wrote to .orchestrator/wes2/.
        # They are allowed to exist but must NOT pollute diff, analyzer, or commit.
        code_files = [f for f in allowed_files if not is_orchestrator_log_file(f)]
        log_files_found = [f for f in allowed_files if is_orchestrator_log_file(f)]
        if log_files_found:
            logger.log(f"  Excluding {len(log_files_found)} orchestrator log file(s) from diff/analyzer/commit.")

        # ---- Dart files (code only, no log files) ----
        dart_files = [f for f in code_files if f.endswith(".dart")]

        # ---- Analyzer ----
        logger.log("  Running Dart analyzer...")
        analyzer_ok, analyzer_output = run_dart_analyzer(dart_files, log_dir)
        logger.log(f"  Analyzer result: {'OK' if analyzer_ok else 'ERRORS'}")

        # ---- Diff (code files only, no log files) ----
        diff_text = get_diff_text(code_files)
        diff_path = log_dir / f"phase_{phase_no}_impl_diff_{impl_attempt}.txt"
        diff_path.write_text(diff_text, encoding="utf-8")

        if not diff_text.strip():
            logger.log("  WARNING: No diff detected. Claude may not have written any files.")

        if not _check_runtime():
            return False

        # ---- Implementation review ----
        logger.log("  Calling OpenAI to review implementation...")
        impl_review_prompt = build_implementation_review_prompt(
            phase_no, phase_name, spec_excerpt,
            diff_text, analyzer_output,
            out_of_scope_files, blocked_files,
            previous_feedback=previous_impl_feedback,
        )
        impl_review_label = f"phase_{phase_no}_impl_review_{impl_attempt}"

        try:
            impl_review = call_openai_reviewer(
                impl_review_prompt, args.openai_model, log_dir, impl_review_label
            )
        except RuntimeError as e:
            logger.log(f"  ERROR: OpenAI reviewer failed: {e}")
            return False

        if not impl_review.get("approved", False):
            impl_review = require_spec_refs_review(
                impl_review, args.openai_model, impl_review_prompt,
                log_dir, impl_review_label
            )

        logger.log(f"  Implementation review: approved={impl_review.get('approved')}")
        logger.log(f"  Reason: {impl_review.get('reason', '')}")

        if impl_review.get("approved", False):
            logger.log("  Implementation APPROVED.")

            # ---- Auto-commit (code files only, no log files) ----
            is_hint_phase = (phase_no == 21)
            should_commit = (
                args.auto_commit
                and (not is_hint_phase or args.auto_commit_hints)
                and bool(code_files)
            )

            if should_commit:
                commit_msg = f"WES2 phase {phase_no}: {phase_name}"
                commit_body = (
                    f"Phase: {phase_no} — {phase_name}\n"
                    f"Reviewer summary: {impl_review.get('reason', '')}\n"
                    f"Files changed: {', '.join(code_files)}\n"
                    f"Generated by WES2 orchestrator."
                )
                commit_hash = git_commit(code_files, commit_msg, commit_body)
                if commit_hash:
                    logger.log(f"  Committed: {commit_hash}")
                else:
                    logger.log("  WARNING: Commit produced no hash (nothing to commit?).")
            elif is_hint_phase and not args.auto_commit_hints:
                logger.log(
                    f"  Phase {phase_no} (hint phase) NOT auto-committed. "
                    "Pass --auto-commit-hints to commit after manual testing."
                )
            elif not code_files:
                logger.log("  No code files changed — nothing to commit.")

            return True

        else:
            corrections = impl_review.get("corrections", [])
            spec_refs = impl_review.get("spec_references_used", [])
            logger.log(f"  Implementation REJECTED. Corrections: {corrections}")
            logger.log(f"  Spec refs: {spec_refs}")
            previous_impl_feedback = (
                f"Review JSON:\n{json.dumps(impl_review, indent=2)}\n\n"
                f"You MUST address all corrections. "
                f"Analyzer output:\n{analyzer_output[:1000]}"
            )

    logger.log(f"\n  FAILED: Implementation not approved after {args.max_implementation_attempts} attempts.")
    return False


def _is_blocked(text: str) -> bool:
    """
    Check if Claude output contains an explicit BLOCKED signal.
    Only triggers when a line starts with 'BLOCKED:' (optionally preceded by whitespace).
    Does NOT trigger on phrases like 'No BLOCKED conditions found.'
    """
    return bool(re.search(r"(?im)^\s*BLOCKED\s*:", text))


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def parse_args():
    parser = argparse.ArgumentParser(
        description="WES2 Orchestrator — Claude Code + OpenAI reviewer"
    )
    parser.add_argument("--spec", required=True, help="Path to WES2 spec .docx file")
    parser.add_argument("--start-phase", type=int, default=18, help="Phase to start from (default: 18)")
    parser.add_argument("--max-phases", type=int, default=1, help="Max number of APPROVED phases to complete (default: 1)")
    parser.add_argument("--max-plan-attempts", type=int, default=6, help="Max plan attempts per phase (default: 6)")
    parser.add_argument("--max-implementation-attempts", type=int, default=6, help="Max implementation attempts per phase (default: 6)")
    parser.add_argument(
        "--claude-cmd",
        default=os.environ.get("CLAUDE_CMD", "claude -p"),
        help="Claude CLI command (default: claude -p or $CLAUDE_CMD)",
    )
    parser.add_argument(
        "--claude-timeout-seconds",
        type=int,
        default=2400,
        help="Per-call timeout for Claude subprocess in seconds (default: 2400)",
    )
    parser.add_argument(
        "--max-runtime-seconds",
        type=int,
        default=14400,
        help="Total orchestrator runtime limit in seconds (default: 14400 = 4 hours)",
    )
    parser.add_argument(
        "--openai-model",
        default=os.environ.get("OPENAI_REVIEW_MODEL", "gpt-5.4-mini"),
        help="OpenAI model for reviewer (default: gpt-5.4-mini or $OPENAI_REVIEW_MODEL)",
    )
    parser.add_argument(
        "--revert-out-of-scope", dest="revert_out_of_scope",
        action="store_true", default=True,
        help="Revert out-of-scope file writes (default: True)",
    )
    parser.add_argument(
        "--no-revert-out-of-scope", dest="revert_out_of_scope",
        action="store_false",
    )
    parser.add_argument(
        "--auto-commit", dest="auto_commit",
        action="store_true", default=True,
        help="Auto-commit approved non-hint phases (default: True)",
    )
    parser.add_argument(
        "--no-auto-commit", dest="auto_commit",
        action="store_false",
    )
    parser.add_argument(
        "--auto-commit-hints", dest="auto_commit_hints",
        action="store_true", default=False,
        help="Also auto-commit Phase 21 (hint phase). Default: False — leave uncommitted for manual testing.",
    )
    return parser.parse_args()


def main():
    args = parse_args()

    # Validate environment
    if not os.environ.get("OPENAI_API_KEY"):
        print("ERROR: OPENAI_API_KEY environment variable not set.", file=sys.stderr)
        sys.exit(1)

    if not os.path.isfile(args.spec):
        print(f"ERROR: Spec file not found: {args.spec}", file=sys.stderr)
        sys.exit(1)

    # Change to repo root
    repo_root = get_repo_root()
    os.chdir(repo_root)
    print(f"Repo root: {repo_root}")

    # Require clean git tree
    require_clean_tree()

    # Set up logging
    log_dir = setup_log_dir()
    logger = Logger(log_dir)
    logger.log(f"WES2 Orchestrator starting")
    logger.log(f"Log directory: {log_dir}")
    logger.log(f"Spec: {args.spec}")
    logger.log(f"Start phase: {args.start_phase}")
    logger.log(f"Max phases (approved): {args.max_phases}")
    logger.log(f"Claude cmd: {args.claude_cmd}")
    logger.log(f"Claude timeout per call: {args.claude_timeout_seconds}s")
    logger.log(f"Max total runtime: {args.max_runtime_seconds}s ({args.max_runtime_seconds // 3600}h)")
    logger.log(f"OpenAI model: {args.openai_model}")
    logger.log(f"Auto-commit: {args.auto_commit}")
    logger.log(f"Auto-commit hints: {args.auto_commit_hints}")

    # Read spec
    logger.log("Reading spec document...")
    try:
        full_spec = read_spec_docx(args.spec)
    except Exception as e:
        logger.log(f"ERROR reading spec: {e}")
        logger.close()
        sys.exit(1)

    spec_brief_path = log_dir / "spec_brief.txt"
    spec_brief_path.write_text(COMPACT_SPEC_BRIEF, encoding="utf-8")
    logger.log(f"Spec read: {len(full_spec)} chars")

    # Determine phase sequence
    all_phases = sorted(PHASE_MAP.keys())
    phases_to_run = [p for p in all_phases if p >= args.start_phase]
    if not phases_to_run:
        # Allow arbitrary phase numbers not in PHASE_MAP
        phases_to_run = list(range(args.start_phase, args.start_phase + args.max_phases + 10))

    approved_count = 0
    exit_code = 0
    start_time = time.monotonic()

    for phase_no in phases_to_run:
        if approved_count >= args.max_phases:
            logger.log(f"\nReached max-phases ({args.max_phases} approved). Stopping.")
            break

        # Total runtime guard before starting a new phase
        elapsed = time.monotonic() - start_time
        if elapsed >= args.max_runtime_seconds:
            logger.log(
                f"\nTIMEOUT: Total runtime {elapsed:.0f}s exceeded "
                f"--max-runtime-seconds {args.max_runtime_seconds}. Stopping before phase {phase_no}."
            )
            exit_code = 1
            break

        phase_name = PHASE_MAP.get(phase_no, f"Phase {phase_no}")
        success = run_phase(phase_no, full_spec, args, log_dir, logger, start_time=start_time)

        if success:
            approved_count += 1
            logger.log(f"\nPhase {phase_no} ({phase_name}) completed successfully. "
                      f"[{approved_count}/{args.max_phases} approved phases done]")
        else:
            # Check if it was a BLOCKED stop
            logger.log(f"\nPhase {phase_no} ({phase_name}) FAILED or BLOCKED.")
            logger.log(f"Log directory: {log_dir}")
            exit_code = 1
            break

    logger.log(f"\n{'='*60}")
    logger.log(f"ORCHESTRATOR COMPLETE")
    logger.log(f"Approved phases: {approved_count}/{args.max_phases}")
    logger.log(f"Log directory: {log_dir}")
    logger.log(f"{'='*60}")
    logger.close()

    sys.exit(exit_code)


if __name__ == "__main__":
    main()


