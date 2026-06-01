#!/usr/bin/env python3
"""
WES2 reviewer/implementer orchestrator.

Purpose:
- Claude Code implements WES2 in small phases.
- OpenAI reviews each phase against the WES2 spec.
- The script enforces a hard write scope so current WES/BB3/PMU/progression files are protected.

Run from repo root.

Allowed writes:
- New WES2 files whose file names begin with WES2_:
  - lib/WES2_*.dart
  - lib/WES2_widgets/WES2_*.dart
  - test/wes2/WES2_*.dart
- docs/wes2/**
- .orchestrator/wes2/**
- lib/home_screen.dart ONLY for the WES2 quick-access card/import

Blocked:
- lib/workout_entry_screen.dart
- lib/main.dart
- lib/bb3_*.dart
- lib/periodization_model_utils.dart
- lib/progression_engine.dart
- pubspec.yaml
- assets/**
- generated files
"""

from __future__ import annotations

import argparse
import fnmatch
import json
import os
import shlex
import subprocess
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

try:
    from openai import OpenAI
except Exception:
    OpenAI = None


ALLOWED_PREFIXES = [
    "docs/wes2/",
    ".orchestrator/wes2/",
]

ALLOWED_GLOBS = [
    "lib/WES2_*.dart",
    "lib/WES2_widgets/WES2_*.dart",
    "test/wes2/WES2_*.dart",
]

# Existing-file exception. Reviewer prompt further restricts this to quick-access navigation only.
ALLOWED_EXACT = [
    "lib/home_screen.dart",
]

BLOCKED_GLOBS = [
    "lib/workout_entry_screen.dart",
    "lib/main.dart",
    "lib/bb3_*.dart",
    "lib/periodization_model_utils.dart",
    "lib/progression_engine.dart",
    "pubspec.yaml",
    "assets/**",
    "**/*.g.dart",
    "**/*.freezed.dart",
    "**/*.mocks.dart",
]


@dataclass
class CmdResult:
    code: int
    out: str


def run(cmd: list[str] | str, *, timeout: int = 1200, input_text: str | None = None) -> CmdResult:
    shell = isinstance(cmd, str)
    p = subprocess.run(
        cmd,
        input=input_text,
        text=True,
        shell=shell,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        timeout=timeout,
    )
    return CmdResult(p.returncode, p.stdout)


def die(msg: str, code: int = 1) -> None:
    print(f"\n❌ {msg}\n")
    sys.exit(code)


def repo_root() -> Path:
    r = run(["git", "rev-parse", "--show-toplevel"], timeout=30)
    if r.code != 0:
        die("This script must be run inside a git repo.")
    return Path(r.out.strip())


def _extract_docx_table(table) -> list[str]:
    rows: list[str] = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return rows


def read_spec(path: Path) -> str:
    """Read text/md/docx spec. For docx, include paragraphs AND table cell text."""
    if not path.exists():
        die(f"Spec not found: {path}")

    if path.suffix.lower() == ".docx":
        try:
            import docx  # python-docx
        except Exception:
            die("Reading .docx requires: pip install python-docx")
        d = docx.Document(str(path))
        parts: list[str] = []
        for p in d.paragraphs:
            txt = p.text.strip()
            if txt:
                parts.append(txt)
        for t in d.tables:
            parts.extend(_extract_docx_table(t))
        return "\n".join(parts)

    return path.read_text(encoding="utf-8", errors="replace")


def ensure_clean_or_confirm() -> None:
    r = run(["git", "status", "--porcelain"], timeout=30)
    if r.code != 0:
        die("Could not read git status.")
    if r.out.strip():
        die("Working tree is not clean. Commit/stash current work first so scope enforcement is safe.")


def changed_files() -> list[str]:
    r = run(["git", "diff", "--name-only"], timeout=60)
    if r.code != 0:
        die("Could not read git diff.")
    files = [x.strip().replace("\\", "/") for x in r.out.splitlines() if x.strip()]

    r2 = run(["git", "ls-files", "--others", "--exclude-standard"], timeout=60)
    if r2.code == 0:
        files += [x.strip().replace("\\", "/") for x in r2.out.splitlines() if x.strip()]

    return sorted(set(files))


def matches_any(path: str, patterns: Iterable[str]) -> bool:
    p = path.replace("\\", "/")
    return any(fnmatch.fnmatch(p, pat) for pat in patterns)


def allowed(path: str) -> bool:
    p = path.replace("\\", "/")
    if p in ALLOWED_EXACT:
        return True
    if any(p.startswith(pref) for pref in ALLOWED_PREFIXES):
        return True
    if matches_any(p, ALLOWED_GLOBS):
        return True
    return False


def blocked(path: str) -> bool:
    return matches_any(path.replace("\\", "/"), BLOCKED_GLOBS)


def enforce_scope(*, revert: bool) -> tuple[list[str], list[str]]:
    out_of_scope = [p for p in changed_files() if not allowed(p)]
    blocked_touched = [p for p in changed_files() if blocked(p)]
    bad = sorted(set(out_of_scope + blocked_touched))

    if bad and revert:
        for p in bad:
            tracked = run(["git", "ls-files", "--error-unmatch", p], timeout=30)
            if tracked.code == 0:
                run(["git", "checkout", "--", p], timeout=60)
            else:
                pp = Path(p)
                try:
                    if pp.is_file() or pp.is_symlink():
                        pp.unlink()
                except FileNotFoundError:
                    pass
    return out_of_scope, blocked_touched


def get_allowed_changed_files() -> list[str]:
    return [p for p in changed_files() if allowed(p)]


def git_diff_allowed() -> str:
    files = get_allowed_changed_files()
    if not files:
        return ""

    tracked_files = [p for p in files if run(["git", "ls-files", "--error-unmatch", p], timeout=10).code == 0]
    untracked = [p for p in files if p not in tracked_files]
    chunks: list[str] = []
    if tracked_files:
        r = run(["git", "diff", "--"] + tracked_files, timeout=300)
        chunks.append(r.out)
    for p in untracked:
        try:
            txt = Path(p).read_text(encoding="utf-8", errors="replace")
            chunks.append(f"\n--- UNTRACKED FILE: {p} ---\n{txt[:30000]}")
        except Exception:
            chunks.append(f"\n--- UNTRACKED FILE: {p} (binary/unreadable) ---\n")
    return "\n".join(chunks)


def write_log(log_dir: Path, name: str, text: str) -> Path:
    log_dir.mkdir(parents=True, exist_ok=True)
    p = log_dir / name
    p.write_text(text, encoding="utf-8")
    return p


def claude(prompt: str, *, timeout: int, claude_cmd: str) -> str:
    cmd = shlex.split(claude_cmd)
    r = run(cmd, timeout=timeout, input_text=prompt)
    if r.code != 0:
        die(f"Claude command failed:\n{r.out}")
    return r.out


def openai_review(prompt: str, *, model: str, max_output_tokens: int = 5000) -> dict:
    if OpenAI is None:
        die("OpenAI SDK missing. Install with: pip install openai")
    client = OpenAI()
    resp = client.responses.create(
        model=model,
        input=prompt,
        text={"format": {"type": "json_object"}},
        max_output_tokens=max_output_tokens,
    )
    text = resp.output_text
    try:
        return json.loads(text)
    except Exception:
        return {"approved": False, "reason": "Reviewer did not return valid JSON.", "raw": text}


def scope_text() -> str:
    return """
Allowed writes:
- lib/WES2_*.dart
- lib/WES2_widgets/WES2_*.dart
- test/wes2/WES2_*.dart
- docs/wes2/**
- .orchestrator/wes2/**
- lib/home_screen.dart ONLY for adding the WES2 Quick Access card/import

All new WES2 Dart file names must begin with WES2_.
Do not create lowercase wes2_ files.
Do not create lib/wes2/ files.
Do not edit lib/main.dart for v1 route wiring.
Home screen WES2 entry must use direct MaterialPageRoute navigation to the WES2 screen, opened to today's date/current date.
Calendar picker taps must continue opening the original WES/WorkoutPage, not WES2.
""".strip()


def plan_prompt(spec: str) -> str:
    return f"""
You are Claude Code implementing WES2 for the GOODLIFT Flutter/Firebase app.

Read the WES2 spec below. DO NOT EDIT FILES YET.
Produce a concise proposed plan only, then STOP.

{scope_text()}

Hard restrictions:
- Existing files may be read but must not be edited, except lib/home_screen.dart for the single WES2 Quick Access card/import.
- Do not edit lib/workout_entry_screen.dart, lib/main.dart, lib/bb3_*.dart, lib/periodization_model_utils.dart, lib/progression_engine.dart, pubspec.yaml, assets, generated files, or Firestore rules.
- Use current WES only as behavior/widget/helper reference, not architecture template.
- Do not rename, move, delete, or modify asset files or asset references.
- Do not broad-scan the repo. Use targeted anchors.

Required plan contents:
1. file map using only WES2_ file names,
2. exact repo anchors to inspect,
3. phase order,
4. Firestore read/write and merge/transaction strategy,
5. Isar/local draft/fast-reopen/offline queue strategy,
6. multi-device listener strategy,
7. BB3 structural sync/update strategy,
8. hint/progression integration strategy,
9. home-screen Quick Access wiring plan,
10. tests/analyzer strategy,
11. risks and how the plan avoids touching blocked files.

WES2 SPEC:
{spec}
"""


def reviewer_plan_prompt(spec: str, plan: str) -> str:
    return f"""
You are the strict WES2 reviewer.

Return JSON only:
{{
  "approved": true/false,
  "reason": "...",
  "required_changes": ["..."],
  "scope_risks": ["..."]
}}

Review Claude's implementation plan against the spec and write scope.

Must enforce:
{scope_text()}

Reject the plan if it proposes editing blocked files, lower-case wes2 file names, changing existing WES/BB3/PMU/progression, changing assets, changing pubspec.yaml, or routing through main.dart.

SPEC:
{spec}

PLAN:
{plan}
"""


def implement_prompt(spec: str, plan: str, prior_feedback: str, phase_no: int) -> str:
    return f"""
You are Claude Code implementing WES2 phase {phase_no}.

Implement ONLY the next smallest safe phase from the approved plan.

{scope_text()}

Hard restrictions:
- You may read existing files, but do not edit them except lib/home_screen.dart for the WES2 Quick Access card/import.
- Do not edit lib/workout_entry_screen.dart, lib/main.dart, lib/bb3_*.dart, lib/periodization_model_utils.dart, lib/progression_engine.dart, pubspec.yaml, assets, generated files, or Firestore rules.
- Keep the phase small. Prefer compiling stubs over giant incomplete wiring.
- Do not rename, move, delete, or modify asset files or asset references.
- If a change requires a blocked file, print BLOCKED with the exact reason instead of editing it.

At the end, print one of:
- PHASE_DONE: <summary>
- ALL_DONE: <summary>
- BLOCKED: <specific missing info or blocked-file need>

Approved plan:
{plan}

Reviewer feedback to respect:
{prior_feedback}

Spec:
{spec}
"""


def reviewer_diff_prompt(spec: str, plan: str, diff: str, analyzer: str, out_of_scope: list[str], blocked_touched: list[str], phase_output: str) -> str:
    return f"""
You are the strict WES2 reviewer.

Return JSON only:
{{
  "approved": true/false,
  "done": true/false,
  "reason": "...",
  "required_changes": ["..."],
  "scope_violation": true/false,
  "next_phase_guidance": "..."
}}

Approve only if:
- changes stay within the allowed WES2 write scope,
- new WES2 Dart file names begin with WES2_,
- lib/home_screen.dart changes, if any, are limited to WES2 Quick Access import/card navigation,
- no blocked file was edited,
- implementation follows the spec,
- analyzer output is acceptable or only contains unrelated pre-existing warnings,
- the phase is small and reviewable.

Allowed write scope:
{scope_text()}

Out-of-scope files detected by script:
{out_of_scope}

Blocked files detected by script:
{blocked_touched}

Claude phase output:
{phase_output}

Analyzer/test output:
{analyzer[-12000:]}

Diff:
{diff[-60000:]}

Spec:
{spec}

Approved plan:
{plan}
"""


def run_analyzer() -> str:
    parts: list[str] = []
    files = [p for p in get_allowed_changed_files() if p.endswith(".dart")]
    if files:
        fmt = run(["dart", "format"] + files, timeout=300)
        parts.append("== dart format ==\n" + fmt.out)
    else:
        parts.append("== dart format ==\n(no changed Dart files)")
    ana = run(["flutter", "analyze"], timeout=1200)
    parts.append("== flutter analyze ==\n" + ana.out)
    return "\n\n".join(parts)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--spec", required=True, help="Path to WES2 spec .md/.txt/.docx")
    ap.add_argument("--max-phases", type=int, default=20)
    ap.add_argument("--claude-cmd", default=os.environ.get("CLAUDE_CMD", "claude -p"))
    ap.add_argument("--openai-model", default=os.environ.get("OPENAI_REVIEW_MODEL"))
    ap.add_argument("--revert-out-of-scope", action=argparse.BooleanOptionalAction, default=True)
    args = ap.parse_args()

    if not args.openai_model:
        die("Set OPENAI_REVIEW_MODEL, e.g. export OPENAI_REVIEW_MODEL='your-review-model'")

    root = repo_root()
    os.chdir(root)
    ensure_clean_or_confirm()

    log_dir = root / ".orchestrator" / "wes2" / time.strftime("%Y%m%d_%H%M%S")
    spec = read_spec(Path(args.spec))
    if len(spec) < 10000:
        die("Spec text looks too short. Check the --spec path or DOCX extraction.")
    write_log(log_dir, "00_spec_snapshot.txt", spec)

    print("→ Asking Claude for plan...")
    plan = claude(plan_prompt(spec), timeout=1800, claude_cmd=args.claude_cmd)
    write_log(log_dir, "01_claude_plan.txt", plan)

    print("→ Reviewing plan with OpenAI...")
    plan_review = openai_review(reviewer_plan_prompt(spec, plan), model=args.openai_model)
    write_log(log_dir, "02_plan_review.json", json.dumps(plan_review, indent=2))

    if not plan_review.get("approved"):
        die("Plan not approved. See log: " + str(log_dir / "02_plan_review.json"))

    feedback = json.dumps(plan_review, indent=2)

    for phase in range(1, args.max_phases + 1):
        print(f"\n→ Claude implementing phase {phase}...")
        out = claude(implement_prompt(spec, plan, feedback, phase), timeout=3600, claude_cmd=args.claude_cmd)
        write_log(log_dir, f"phase_{phase:02d}_claude.txt", out)

        out_of_scope, blocked_touched = enforce_scope(revert=args.revert_out_of_scope)
        if out_of_scope or blocked_touched:
            print("⚠️ Scope issue detected and reverted/removed where possible:")
            for b in sorted(set(out_of_scope + blocked_touched)):
                print(" -", b)

        print("→ Running formatter/analyzer...")
        analyzer = run_analyzer()
        write_log(log_dir, f"phase_{phase:02d}_analyzer.txt", analyzer)

        diff = git_diff_allowed()
        write_log(log_dir, f"phase_{phase:02d}_diff.patch", diff)

        print("→ Reviewing phase with OpenAI...")
        review = openai_review(
            reviewer_diff_prompt(spec, plan, diff, analyzer, out_of_scope, blocked_touched, out),
            model=args.openai_model,
        )
        write_log(log_dir, f"phase_{phase:02d}_review.json", json.dumps(review, indent=2))

        feedback = json.dumps(review, indent=2)

        if not review.get("approved"):
            print("Phase not approved. Claude will receive reviewer feedback next phase.")
            continue

        if review.get("done") or "ALL_DONE" in out:
            print("\n✅ WES2 orchestration complete.")
            print("Logs:", log_dir)
            print("Review final changed files with: git diff --stat")
            return

    die(f"Reached max phases ({args.max_phases}) without completion. Logs: {log_dir}")


if __name__ == "__main__":
    main()
