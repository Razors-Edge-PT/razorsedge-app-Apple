#!/usr/bin/env python3
r"""
WES2 autonomous spec-grounded reviewer/implementer orchestrator.

Claude implements. OpenAI reviews. The script enforces WES2-only writes.

Key behaviour:
- max-phases means number of APPROVED phases completed, not number of attempts.
- Each phase has repeated plan/review and implement/review correction loops.
- Reviewer corrections must cite WES2 spec references.
- Existing non-WES2 files may be read for schema/reference, but not written.
- WES2 Dart implementation files must contain WES2 in the filename/path.
- Approved non-hint phases auto-commit by default with phase-specific messages.
- Hint phase does not auto-commit by default so it can be manually tested first.

Example:
python .\docs\wes2_orchestrator.py --spec ".\docs\wes2\WES2_Product_Spec_and_Implementation_Contract_v1_revised_with_examples.docx" --start-phase 18 --max-phases 3
"""

from __future__ import annotations

import argparse
import fnmatch
import json
import os
import re
import shlex
import subprocess
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

HINT_PHASE_NO = 21

PHASE_MAP: dict[int, dict[str, Any]] = {
    17: {"name": "Day Timer + Workout Summary", "keywords": ["timer", "summary", "workout summary", "floating timer", "sets logged", "Add Circuit"], "commit": "WES2 phase 17: timer and workout summary"},
    18: {"name": "Templates", "keywords": ["template", "templates", "save as template", "load template", "templateLoaded"], "commit": "WES2 phase 18: templates"},
    19: {"name": "Settings Cog Dialog", "keywords": ["settings", "exerciseSettings", "increments", "weekly frequency", "periodization", "RIR", "progression"], "commit": "WES2 phase 19: exercise settings dialog"},
    20: {"name": "BB3 Structural Handoff", "keywords": ["BB3", "planned day", "delete", "replace", "move", "structural", "updatePlannedDay"], "commit": "WES2 phase 20: BB3 structural handoff"},
    21: {"name": "Hint Engine Integration", "keywords": ["hint", "hints", "E1RM", "PMU", "progression", "ProgressionEngine", "modelHint"], "commit": "WES2 phase 21: hint engine integration"},
    22: {"name": "Final WES2 Spec Compliance Polish", "keywords": ["acceptance", "test matrix", "polish", "final", "overflow", "offline", "multi-device"], "commit": "WES2 phase 22: final spec compliance polish"},
}

ALLOWED_PREFIXES = ["docs/wes2/", ".orchestrator/wes2/"]
ALLOWED_EXACT = ["docs/wes2_orchestrator.py"]
ALLOWED_GLOBS = ["lib/WES2_*.dart", "lib/WES2_widgets/WES2_*.dart", "test/wes2/WES2_*.dart"]
BLOCKED_GLOBS = [
    "lib/workout_entry_screen.dart",
    "lib/main.dart",
    "lib/bb3_*.dart",
    "lib/periodization_model_utils.dart",
    "lib/progression_engine.dart",
    "pubspec.yaml",
    "pubspec.lock",
    "assets/**",
    "**/*.g.dart",
    "**/*.freezed.dart",
    "**/*.mocks.dart",
    ".claude/settings.local.json",
]
DO_NOT_AUTOCOMMIT_PREFIXES = [".orchestrator/"]


@dataclass
class CmdResult:
    code: int
    out: str


def run(cmd: list[str] | str, *, timeout: int = 1200, input_text: str | None = None) -> CmdResult:
    shell = isinstance(cmd, str)
    p = subprocess.run(cmd, input=input_text, text=True, shell=shell, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, timeout=timeout)
    return CmdResult(p.returncode, p.stdout)


def die(msg: str, code: int = 1) -> None:
    print(f"\n❌ {msg}\n")
    sys.exit(code)


def repo_root() -> Path:
    r = run(["git", "rev-parse", "--show-toplevel"], timeout=30)
    if r.code != 0:
        die("This script must be run inside a git repo.")
    return Path(r.out.strip())


def _extract_docx_table(table) -> list[str]:
    rows: list[str] = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return rows


def read_spec(path: Path) -> str:
    if not path.exists():
        die(f"Spec not found: {path}")
    if path.suffix.lower() == ".docx":
        try:
            import docx
        except Exception:
            die("Reading .docx requires: pip install python-docx")
        d = docx.Document(str(path))
        parts: list[str] = []
        for p in d.paragraphs:
            txt = p.text.strip()
            if txt:
                parts.append(txt)
        for t in d.tables:
            parts.extend(_extract_docx_table(t))
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


def ensure_clean_or_confirm() -> None:
    r = run(["git", "status", "--porcelain"], timeout=30)
    if r.code != 0:
        die("Could not read git status.")
    dirty = []
    for line in r.out.splitlines():
        path = line[3:].replace("\\", "/") if len(line) > 3 else ""
        if path.startswith(".orchestrator/wes2/"):
            continue
        dirty.append(line)
    if dirty:
        die("Working tree is not clean. Commit/stash current work first so scope enforcement is safe.\n" + "\n".join(dirty))


def changed_files(include_logs: bool = True) -> list[str]:
    r = run(["git", "diff", "--name-only"], timeout=60)
    if r.code != 0:
        die("Could not read git diff.")
    files = [x.strip().replace("\\", "/") for x in r.out.splitlines() if x.strip()]
    r2 = run(["git", "ls-files", "--others", "--exclude-standard"], timeout=60)
    if r2.code == 0:
        files += [x.strip().replace("\\", "/") for x in r2.out.splitlines() if x.strip()]
    files = sorted(set(files))
    if not include_logs:
        files = [p for p in files if not any(p.startswith(pref) for pref in DO_NOT_AUTOCOMMIT_PREFIXES)]
    return files


def matches_any(path: str, patterns: Iterable[str]) -> bool:
    p = path.replace("\\", "/")
    return any(fnmatch.fnmatch(p, pat) for pat in patterns)


def allowed(path: str) -> bool:
    p = path.replace("\\", "/")
    if p in ALLOWED_EXACT:
        return True
    if any(p.startswith(pref) for pref in ALLOWED_PREFIXES):
        return True
    if matches_any(p, ALLOWED_GLOBS):
        return True
    return False


def blocked(path: str) -> bool:
    return matches_any(path.replace("\\", "/"), BLOCKED_GLOBS)


def enforce_scope(*, revert: bool) -> tuple[list[str], list[str]]:
    cf = changed_files(include_logs=True)
    out_of_scope = [p for p in cf if not allowed(p)]
    blocked_touched = [p for p in cf if blocked(p)]
    bad = sorted(set(out_of_scope + blocked_touched))
    if bad and revert:
        for p in bad:
            tracked = run(["git", "ls-files", "--error-unmatch", p], timeout=30)
            if tracked.code == 0:
                run(["git", "checkout", "--", p], timeout=60)
            else:
                pp = Path(p)
                try:
                    if pp.is_file() or pp.is_symlink():
                        pp.unlink()
                except FileNotFoundError:
                    pass
    return out_of_scope, blocked_touched


def get_allowed_changed_files(include_logs: bool = False) -> list[str]:
    return [p for p in changed_files(include_logs=include_logs) if allowed(p)]


def git_diff_allowed(max_chars: int = 50000) -> str:
    files = get_allowed_changed_files(include_logs=False)
    if not files:
        return ""
    tracked_files = [p for p in files if run(["git", "ls-files", "--error-unmatch", p], timeout=10).code == 0]
    untracked = [p for p in files if p not in tracked_files]
    chunks: list[str] = []
    if tracked_files:
        r = run(["git", "diff", "--"] + tracked_files, timeout=300)
        chunks.append(r.out)
    for p in untracked:
        try:
            txt = Path(p).read_text(encoding="utf-8", errors="replace")
            chunks.append(f"\n--- UNTRACKED FILE: {p} ---\n{txt[:22000]}")
        except Exception:
            chunks.append(f"\n--- UNTRACKED FILE: {p} (binary/unreadable) ---\n")
    return "\n".join(chunks)[-max_chars:]


def write_log(log_dir: Path, name: str, text: str) -> Path:
    log_dir.mkdir(parents=True, exist_ok=True)
    p = log_dir / name
    p.write_text(text, encoding="utf-8")
    return p


def claude(prompt: str, *, timeout: int, claude_cmd: str) -> str:
    cmd = shlex.split(claude_cmd)
    r = run(cmd, timeout=timeout, input_text=prompt)
    if r.code != 0:
        die(f"Claude command failed:\n{r.out}")
    return r.out


def openai_json(prompt: str, *, model: str, max_output_tokens: int = 2600) -> dict[str, Any]:
    if OpenAI is None:
        die("OpenAI SDK missing. Install with: pip install openai")
    client = OpenAI()
    resp = client.responses.create(model=model, input=prompt, text={"format": {"type": "json_object"}}, max_output_tokens=max_output_tokens)
    text = resp.output_text
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed
        return {"approved": False, "reason": "Reviewer returned JSON but not an object.", "raw": text}
    except Exception:
        return {"approved": False, "reason": "Reviewer did not return valid JSON.", "raw": text}


def scope_text() -> str:
    return """
Allowed WRITES:
- lib/WES2_*.dart
  Examples explicitly allowed: lib/WES2_screen.dart, lib/WES2_controller.dart, lib/WES2_models.dart, lib/WES2_repository.dart, lib/WES2_plan_service.dart, lib/WES2_local_store.dart, lib/WES2_template_service.dart, lib/WES2_sync_service.dart.
- lib/WES2_widgets/WES2_*.dart
  Examples explicitly allowed: lib/WES2_widgets/WES2_set_row.dart, lib/WES2_widgets/WES2_exercise_card.dart, lib/WES2_widgets/WES2_template_picker.dart.
- test/wes2/WES2_*.dart
- docs/wes2/**
- .orchestrator/wes2/**
- docs/wes2_orchestrator.py

Dart implementation files must contain WES2 in the filename/path. Do not create non-WES2 Dart implementation files.

Allowed READ-ONLY inspection:
- Existing non-WES2 files may be read for schema/reference only when needed.
- Reading lib/templates.dart, lib/template_model.dart, current WES, BB3, or other reference files is allowed.
- Writing to those files is NOT allowed unless they match the allowed WES2 write scope above.

Blocked WRITES:
- lib/workout_entry_screen.dart
- lib/main.dart
- lib/bb3_*.dart
- lib/periodization_model_utils.dart
- lib/progression_engine.dart
- pubspec.yaml/pubspec.lock
- assets/**
- generated files
- .claude/settings.local.json

Reviewer must distinguish READS from WRITES:
- Do not reject read-only inspection of reference files.
- Reject any write to blocked/reference files.

Do not rename, move, delete, or modify asset files or asset references.
Do not use broad scans.
Do not use ListTile inside PopupMenuItem.
Avoid Flutter RenderFlex overflow: long Row text must use Expanded/Flexible, maxLines, and ellipsis.
Treat any RenderFlex overflow risk as a blocker.
""".strip()


def build_spec_brief(spec: str) -> str:
    keep_patterns = [
        r"(?is)1\. Core Contract.*?(?=\n2\. Existing|\Z)",
        r"(?is)3\. Firestore and Data Model Contract.*?(?=\n4\. State|\Z)",
        r"(?is)4\. State Ownership.*?(?=\n5\. Opening|\Z)",
        r"(?is)5\. Opening, Loading.*?(?=\n6\. Planned|\Z)",
        r"(?is)9\..*?Exercise, set, circuit, template, undo, and delete/replace behavior.*?(?=\n10\.|\Z)",
        r"(?is)10\..*?UI specification.*?(?=\n11\.|\Z)",
        r"(?is)11\..*?Timed exercises, velocity, notes, video, timer, and summary behavior.*?(?=\n12\.|\Z)",
        r"(?is)12\..*?Multi-device.*?(?=\n13\.|\Z)",
        r"(?is)Acceptance criteria and test matrix.*?\Z",
    ]
    chunks: list[str] = []
    for pat in keep_patterns:
        m = re.search(pat, spec)
        if m:
            chunks.append(m.group(0)[:5200])
    brief = "\n\n--- SPEC BRIEF SECTION ---\n\n".join(chunks)
    if len(brief) < 2000:
        brief = spec[:22000]
    return brief[:28000]


def relevant_spec_excerpt(spec: str, phase_no: int, max_chars: int = 8500) -> str:
    keywords = PHASE_MAP.get(phase_no, {}).get("keywords", [])
    if not keywords:
        return build_spec_brief(spec)[:max_chars]
    lines = spec.splitlines()
    hits: list[int] = []
    for i, line in enumerate(lines):
        low = line.lower()
        if any(k.lower() in low for k in keywords):
            hits.append(i)
    windows: list[str] = []
    seen_ranges: list[tuple[int, int]] = []
    for i in hits[:24]:
        start = max(0, i - 22)
        end = min(len(lines), i + 60)
        if any(not (end < a or start > b) for a, b in seen_ranges):
            continue
        seen_ranges.append((start, end))
        windows.append("\n".join(lines[start:end]))
    txt = "\n\n--- RELEVANT SPEC EXCERPT ---\n\n".join(windows)
    if not txt.strip():
        txt = build_spec_brief(spec)
    return txt[:max_chars]


def phase_guidance(phase_no: int) -> str:
    if phase_no == 17:
        return """
Phase 17 target:
- Day-level floating timer and Workout Summary.
- Timer AppBar overflow item should open/toggle floating timer.
- Templates remains placeholder.
- Summary should be accessible from the bottom/Add Circuit row if feasible because the spec places Summary there.
- Do not implement timed-exercise reps-as-seconds.
- Do not persist timer to Firestore/Isar.
- Do not touch hints/PMU/progression/templates/settings.
""".strip()

    if phase_no == 18:
        return """
Phase 18 target:
- Templates only.
- This is a continuation from the current WES2 repo state, not a WES2 restart.
- Implement the next smallest safe template phase inside WES2-specific files only.

Allowed writes for this phase:
- lib/WES2_*.dart
- lib/WES2_widgets/WES2_*.dart
- test/wes2/WES2_*.dart if tests are added

Read-only reference inspection allowed:
- lib/templates.dart
- lib/template_model.dart
- existing template-related files
These may be read for schema/reference only. Do not edit them.

Required Phase 18 template contract:
1. Template-loaded blank rows:
   - On load, blank template-loaded rows persist in wesPlannedExercises[].
   - Do not imply template-loaded rows are permanently constrained to wesPlannedExercises[].
   - Later completed/logged values must still flow into exercises[] through the existing WES2 save/done pathways.
   - Do not store actual logged values inside saved templates.

2. Duplicate exerciseId prevention:
   - Template load must silently dedupe duplicate exerciseIds across the current date.
   - The same exerciseId must not appear twice across circuits.
   - If duplicate template rows are encountered, keep the first valid row in template/global order and skip later duplicates.
   - Do not broad-scan the repo to enforce this.

3. Template load replacement semantics:
   - Loading a template replaces the current WES2 day/session rows according to the WES2 template contract.
   - Snapshot undo before replacing current WES2 rows.
   - Preserve current selected date/session behaviour.
   - Preserve local draft/fast reopen behaviour.
   - Preserve expanded/collapsed and scroll/edit-anchor behaviour if already present.
   - Do not reset or break Add Exercise/Add Circuit/Add Set/Remove Set/Delete/Replace/Move/Done behaviour.

4. BB3 planned-day handling:
   - If the current WES2 day contains BB3-sourced rows and the template is explicitly replacing that day, BB3 clearing may be planned only for the current selected BB3 planned day.
   - BB3 clearing must be tightly limited to the current blockId/weekIndex/dayIndex.
   - It must not touch non-current BB3 days.
   - It must not touch completed workout documents.
   - It must not modify BB3 files.
   - If the BB3 current-day write shape is not fully verified from WES2 services, print BLOCKED rather than guessing.

5. Save workout as template:
   - Save structure only: exerciseId, name, circuitIndex, orderIndex, setCount, and template metadata as supported by existing schema.
   - Do not save actual values, hints, execution notes, plan notes, isMarkedDone, or completed workout reality.
   - Eligibility must follow spec. If eligibility is uncertain, gate conservatively or print BLOCKED.

6. UI/layout:
   - Template picker/dialog/bottom sheet must use WES2-specific widgets/files only.
   - All long text in Rows must be wrapped in Expanded/Flexible.
   - Use maxLines and TextOverflow.ellipsis.
   - Do not use ListTile inside PopupMenuItem.
   - RenderFlex overflow is a blocker.

7. Non-goals:
   - Do not implement hints/PMU/progression.
   - Do not implement settings dialog.
   - Do not implement BB3 structural handoff beyond the tightly scoped template replacement clear described above.
   - Do not edit current WES, BB3 files, PMU, progression, assets, pubspec, generated files, or .claude/settings.local.json.

The Phase 18 plan must explicitly mention:
- continuation from current repo state,
- WES2-only write scope,
- read-only reference inspection only,
- template-loaded blank rows start in wesPlannedExercises[] but can later flow to exercises[] through normal save,
- duplicate exerciseId dedupe,
- undo snapshot before replacement,
- current-day-only BB3 clear limits,
- overflow safeguards.
""".strip()

    if phase_no == 19:
        return """
Phase 19 target:
- Settings cog dialog.
- Must use exerciseSettings only; do not write plannedExerciseDetails.
- Requires activeBlockId guard.
- Internet-required; no offline queue.
- Do not implement hint recalculation beyond reload/refresh trigger.
""".strip()

    if phase_no == 20:
        return """
Phase 20 target:
- BB3 structural handoff for BB3-sourced delete/replace/move/reorder.
- Read planned day write path carefully.
- Do not mutate BB3 until path and row shape are verified.
""".strip()

    if phase_no == 21:
        return """
Phase 21 target:
- Hint engine integration.
- Must respect BB3 planned overrides as hints only, never actual values.
- Do not overwrite typed WES2 values.
""".strip()

    return "Continue the next smallest safe WES2 phase from the approved plan."



def initial_plan_prompt(spec_brief: str, phase_no: int, prior_feedback: str = "") -> str:
    return f"""
You are Claude Code planning WES2 phase {phase_no}: {PHASE_MAP.get(phase_no, {}).get("name", "next phase")}.

Do NOT edit files.
Produce a concise implementation plan only, then STOP.

This is a continuation from the current repo state. Do not restart WES2 architecture.

{scope_text()}

Phase-specific guidance:
{phase_guidance(phase_no)}

Required plan:
1. Confirm this continues from the current repo state and is not a WES2 restart.
2. Confirm current phase target and that only this incremental phase will be implemented.
3. Files/anchors to inspect, separating READ-ONLY reference inspection from WRITABLE WES2 files.
4. Small implementation steps.
5. Spec rules you will follow by section/topic name.
6. Explicit continuity protections:
   - preserve selected date/session behaviour,
   - preserve local draft/fast reopen behaviour,
   - preserve undo stack continuity,
   - preserve Add Exercise/Add Circuit/Add Set/Remove Set/Delete/Replace/Move/Done behaviour,
   - preserve duplicate exerciseId prevention across the date,
   - preserve no-current-WES / no-BB3 / no-PMU / no-progression modifications unless this phase explicitly permits WES2-scoped service interaction.
7. If this is Phase 18 Templates, explicitly state:
   - template load is a WES2-only extension, not an architecture restart,
   - template load snapshots undo before replacing current WES2 rows,
   - template load dedupes duplicate exerciseIds,
   - template load replaces current WES2 day rows only as specified by the WES2 template contract,
   - save-to-template does not store actual logged values,
   - existing template files may be read for schema only,
   - blocked/template reference files are not edited.
8. Firestore/local persistence impact.
9. Analyzer/static-test command.
10. Risks and overflow guards.

Reviewer feedback to address if any:
{prior_feedback}

COMPACT SPEC BRIEF:
{spec_brief}
"""


def reviewer_plan_prompt(spec_brief: str, plan: str, phase_no: int) -> str:
    return f"""
You are the strict WES2 reviewer.

Return JSON only:
{{
  "approved": true/false,
  "reason": "...",
  "required_changes": ["..."],
  "spec_references_used": ["..."],
  "scope_risks": ["..."]
}}

Review Claude's plan for phase {phase_no}: {PHASE_MAP.get(phase_no, {}).get("name", "next phase")}.

Approve if the plan is a safe, WES2-scoped continuation that can implement this phase while preserving existing WES2 behaviour.

Must enforce:
{scope_text()}

Reviewer rules:
- If you reject or request a correction, "spec_references_used" MUST be non-empty and cite/spec-reference the supporting WES2 spec topic/section.
- Do not reject read-only inspection of non-WES2 reference files.
- Reject writes to blocked/reference files.
- Do not require the plan to prove implementation details that are better verified during diff review, unless the omission creates a clear spec/scope risk.
- Do not reject merely because the phase uses existing WES2 files such as lib/WES2_screen.dart, lib/WES2_controller.dart, or lib/WES2_repository.dart; those are allowed writes.

COMPACT SPEC BRIEF:
{spec_brief}

PLAN:
{plan}
"""


def revise_plan_prompt(plan: str, review: dict[str, Any], phase_no: int, spec_brief: str) -> str:
    return f"""
Your WES2 phase {phase_no} plan was rejected by the reviewer.

Do NOT edit files.
Revise the plan only, addressing every required change.
Keep it concise and scoped to phase {phase_no}: {PHASE_MAP.get(phase_no, {}).get("name", "next phase")}.

Reviewer JSON:
{json.dumps(review, indent=2)}

Original plan:
{plan}

Scope:
{scope_text()}

Phase guidance:
{phase_guidance(phase_no)}

Relevant compact spec brief:
{spec_brief[:14000]}
"""


def implement_prompt(spec_excerpt: str, plan: str, feedback: str, phase_no: int, attempt: int) -> str:
    return f"""
You are Claude Code implementing WES2 phase {phase_no}: {PHASE_MAP.get(phase_no, {}).get("name", "next phase")}.

Implementation attempt {attempt}.
Implement ONLY this phase.
Keep changes small and reviewable.

{scope_text()}

Phase-specific guidance:
{phase_guidance(phase_no)}

Hard restrictions:
- Do not edit blocked files.
- If a change requires a blocked file, print BLOCKED with the exact reason.
- Do not touch hints/PMU/progression unless this is explicitly phase 21.
- Do not rename, move, delete, or modify asset files or asset references.
- Avoid RenderFlex overflows.
- Do not use ListTile inside PopupMenuItem.

At the end, print:
PHASE_DONE: <summary>
or
BLOCKED: <reason>

Reviewer feedback to respect:
{feedback}

Approved phase plan:
{plan[:10000]}

Relevant spec excerpt for this phase:
{spec_excerpt}
"""


def correction_prompt(spec_excerpt: str, plan: str, review: dict[str, Any], analyzer: str, diff: str, phase_no: int, attempt: int) -> str:
    return f"""
Your WES2 phase {phase_no} implementation was not approved.

Correction attempt {attempt}.
Make ONLY the required corrections from the reviewer.
Do not broaden scope.
Do not restart the phase.
Do not touch blocked files.
Do not touch unrelated files.

Reviewer JSON:
{json.dumps(review, indent=2)}

Recent analyzer output:
{analyzer[-7000:]}

Current allowed diff:
{diff[-30000:]}

Scope:
{scope_text()}

Approved phase plan:
{plan[:8000]}

Relevant spec excerpt:
{spec_excerpt}

At the end, print:
PHASE_DONE: <correction summary>
or
BLOCKED: <reason>
"""


def reviewer_diff_prompt(spec_excerpt: str, plan: str, diff: str, analyzer: str, out_of_scope: list[str], blocked_touched: list[str], phase_output: str, phase_no: int) -> str:
    return f"""
You are the strict WES2 reviewer.

Return JSON only:
{{
  "approved": true/false,
  "done": true/false,
  "reason": "...",
  "required_changes": ["..."],
  "spec_references_used": ["..."],
  "scope_violation": true/false,
  "next_phase_guidance": "..."
}}

Approve only if:
- changes stay within allowed WES2 write scope,
- no blocked file was edited,
- read-only inspection of reference files is not treated as a scope violation,
- implementation follows the relevant spec excerpt,
- analyzer output is acceptable or only unrelated/pre-existing warnings,
- no RenderFlex/overflow-prone layout was introduced,
- the phase remains small and reviewable.

Reviewer rules:
- If you reject or request a correction, "spec_references_used" MUST be non-empty and cite/spec-reference the supporting WES2 spec topic/section.
- Do not reject merely because implementation uses existing WES2 files such as lib/WES2_screen.dart, lib/WES2_controller.dart, or lib/WES2_repository.dart.
- Reject any write to blocked/reference files.
- For hint phase, be extra strict and do not allow commits unless human approval follows.

Phase {phase_no}: {PHASE_MAP.get(phase_no, {}).get("name", "next phase")}

Allowed write scope:
{scope_text()}

Out-of-scope files detected by script:
{out_of_scope}

Blocked files detected by script:
{blocked_touched}

Claude phase output:
{phase_output[-9000:]}

Analyzer/static output:
{analyzer[-8000:]}

Allowed diff:
{diff[-42000:]}

Relevant spec excerpt:
{spec_excerpt}

Approved phase plan:
{plan[:6000]}
"""


def run_static_checks() -> str:
    parts: list[str] = []
    files = [p for p in get_allowed_changed_files(include_logs=False) if p.endswith(".dart")]
    if files:
        fmt = run(["dart", "format"] + files, timeout=300)
        parts.append("== dart format changed WES2 files ==\n" + fmt.out)
        ana = run(["flutter", "analyze"] + files, timeout=1200)
        parts.append("== targeted flutter analyze changed WES2 files ==\n" + ana.out)
    else:
        parts.append("== dart format ==\n(no changed Dart files)")
        parts.append("== targeted flutter analyze ==\n(no changed Dart files)")
    diff_check = run(["git", "diff", "--check"], timeout=120)
    parts.append("== git diff --check ==\n" + diff_check.out)
    test_dir = Path("test/wes2")
    if test_dir.exists() and list(test_dir.glob("WES2_*.dart")):
        tests = run(["flutter", "test", "test/wes2"], timeout=1800)
        parts.append("== flutter test test/wes2 ==\n" + tests.out)
    return "\n\n".join(parts)


def changed_files_for_commit() -> list[str]:
    files = [p for p in get_allowed_changed_files(include_logs=False)]
    files = [p for p in files if not any(p.startswith(pref) for pref in DO_NOT_AUTOCOMMIT_PREFIXES)]
    files = [p for p in files if p != "docs/wes2_orchestrator.py"]
    return files


def auto_commit_phase(phase_no: int, review: dict[str, Any], claude_out: str) -> bool:
    files = changed_files_for_commit()
    if not files:
        print("No implementation files to auto-commit.")
        return False
    add = run(["git", "add"] + files, timeout=120)
    if add.code != 0:
        die("git add failed:\n" + add.out)
    summary = str(review.get("reason") or "").strip() or (str(claude_out).strip().splitlines()[-1][:300] if claude_out.strip() else "Approved WES2 phase changes")
    msg = PHASE_MAP.get(phase_no, {}).get("commit", f"WES2 phase {phase_no}: approved changes")
    body = [
        f"Phase {phase_no}: {PHASE_MAP.get(phase_no, {}).get('name', 'WES2 phase')}",
        "",
        "Approved by OpenAI reviewer against WES2 spec.",
        "",
        "Reviewer summary:",
        summary[:1200],
        "",
        "Files changed:",
        *[f"- {p}" for p in files],
    ]
    commit = run(["git", "commit", "-m", msg, "-m", "\n".join(body)], timeout=300)
    if commit.code != 0:
        die("git commit failed:\n" + commit.out)
    print(commit.out)
    return True


def approve_plan_loop(*, spec_brief: str, phase_no: int, openai_model: str, claude_cmd: str, max_attempts: int, log_dir: Path) -> str:
    plan = ""
    feedback_json = "{}"
    for attempt in range(1, max_attempts + 1):
        print(f"→ Planning phase {phase_no}, attempt {attempt}/{max_attempts}...")
        if attempt == 1:
            prompt = initial_plan_prompt(spec_brief, phase_no, feedback_json)
        else:
            prompt = revise_plan_prompt(plan, json.loads(feedback_json), phase_no, spec_brief)
        plan = claude(prompt, timeout=1800, claude_cmd=claude_cmd)
        write_log(log_dir, f"phase_{phase_no:02d}_plan_attempt_{attempt}.txt", plan)
        print("→ Reviewing plan with OpenAI...")
        review = require_spec_refs_review(reviewer_plan_prompt(spec_brief, plan, phase_no), model=openai_model, log_dir=log_dir, log_name_prefix=f"phase_{phase_no:02d}_plan_review_attempt_{attempt}")
        write_log(log_dir, f"phase_{phase_no:02d}_plan_review_attempt_{attempt}.json", json.dumps(review, indent=2))
        feedback_json = json.dumps(review, indent=2)
        if review.get("approved"):
            print(f"✅ Phase {phase_no} plan approved.")
            return plan
        print("Plan not approved; sending reviewer feedback back to Claude.")
    die(f"Phase {phase_no} plan failed after {max_attempts} attempts. Logs: {log_dir}")


def implementation_loop(*, spec: str, plan: str, phase_no: int, openai_model: str, claude_cmd: str, max_attempts: int, log_dir: Path, revert_out_of_scope: bool) -> dict[str, Any]:
    spec_excerpt = relevant_spec_excerpt(spec, phase_no)
    write_log(log_dir, f"phase_{phase_no:02d}_spec_excerpt.txt", spec_excerpt)
    feedback = "{}"
    analyzer = ""
    diff = ""
    last_out = ""
    for attempt in range(1, max_attempts + 1):
        print(f"\n→ Implementing phase {phase_no}, attempt {attempt}/{max_attempts}...")
        if attempt == 1:
            prompt = implement_prompt(spec_excerpt, plan, feedback, phase_no, attempt)
        else:
            prompt = correction_prompt(spec_excerpt, plan, json.loads(feedback), analyzer, diff, phase_no, attempt)
        last_out = claude(prompt, timeout=3600, claude_cmd=claude_cmd)
        write_log(log_dir, f"phase_{phase_no:02d}_impl_attempt_{attempt}_claude.txt", last_out)
        if "BLOCKED:" in last_out:
            die(f"Claude reported BLOCKED in phase {phase_no}. See log: {log_dir}")
        out_of_scope, blocked_touched = enforce_scope(revert=revert_out_of_scope)
        if out_of_scope or blocked_touched:
            print("⚠️ Scope issue detected and reverted/removed where possible:")
            for b in sorted(set(out_of_scope + blocked_touched)):
                print(" -", b)
        print("→ Running formatter/static checks...")
        analyzer = run_static_checks()
        write_log(log_dir, f"phase_{phase_no:02d}_impl_attempt_{attempt}_static.txt", analyzer)
        diff = git_diff_allowed()
        write_log(log_dir, f"phase_{phase_no:02d}_impl_attempt_{attempt}_diff.patch", diff)
        print("→ Reviewing implementation with OpenAI...")
        review = require_spec_refs_review(reviewer_diff_prompt(spec_excerpt, plan, diff, analyzer, out_of_scope, blocked_touched, last_out, phase_no), model=openai_model, log_dir=log_dir, log_name_prefix=f"phase_{phase_no:02d}_impl_review_attempt_{attempt}")
        write_log(log_dir, f"phase_{phase_no:02d}_impl_review_attempt_{attempt}.json", json.dumps(review, indent=2))
        feedback = json.dumps(review, indent=2)
        if review.get("approved"):
            print(f"✅ Phase {phase_no} implementation approved.")
            review["_claude_out"] = last_out
            return review
        print("Implementation not approved; sending reviewer feedback back to Claude for correction.")
    die(f"Phase {phase_no} implementation failed after {max_attempts} attempts. Logs: {log_dir}")


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--spec", required=True)
    ap.add_argument("--start-phase", type=int, default=18)
    ap.add_argument("--max-phases", type=int, default=1)
    ap.add_argument("--max-plan-attempts", type=int, default=6)
    ap.add_argument("--max-implementation-attempts", type=int, default=6)
    ap.add_argument("--claude-cmd", default=os.environ.get("CLAUDE_CMD", "claude -p"))
    ap.add_argument("--openai-model", default=os.environ.get("OPENAI_REVIEW_MODEL"))
    ap.add_argument("--revert-out-of-scope", action=argparse.BooleanOptionalAction, default=True)
    ap.add_argument("--auto-commit", action=argparse.BooleanOptionalAction, default=True)
    ap.add_argument("--auto-commit-hints", action=argparse.BooleanOptionalAction, default=False)
    args = ap.parse_args()
    if not args.openai_model:
        die("Set OPENAI_REVIEW_MODEL, e.g. $env:OPENAI_REVIEW_MODEL='gpt-5.4-mini'")
    root = repo_root()
    os.chdir(root)
    ensure_clean_or_confirm()
    log_dir = root / ".orchestrator" / "wes2" / time.strftime("%Y%m%d_%H%M%S")
    spec = read_spec(Path(args.spec))
    if len(spec) < 10000:
        die("Spec text looks too short. Check the --spec path or DOCX extraction.")
    spec_brief = build_spec_brief(spec)
    write_log(log_dir, "00_spec_brief.txt", spec_brief)
    phase_no = args.start_phase
    completed = 0
    while completed < args.max_phases:
        print("\n══════════════════════════════════════════════════════")
        print(f"Phase {phase_no}: {PHASE_MAP.get(phase_no, {}).get('name', 'WES2 phase')}")
        print("══════════════════════════════════════════════════════")
        plan = approve_plan_loop(spec_brief=spec_brief, phase_no=phase_no, openai_model=args.openai_model, claude_cmd=args.claude_cmd, max_attempts=args.max_plan_attempts, log_dir=log_dir)
        review = implementation_loop(spec=spec, plan=plan, phase_no=phase_no, openai_model=args.openai_model, claude_cmd=args.claude_cmd, max_attempts=args.max_implementation_attempts, log_dir=log_dir, revert_out_of_scope=args.revert_out_of_scope)
        should_auto_commit = args.auto_commit
        if phase_no >= HINT_PHASE_NO and not args.auto_commit_hints:
            should_auto_commit = False
        if should_auto_commit:
            auto_commit_phase(phase_no, review, str(review.get("_claude_out", "")))
        else:
            print(f"Auto-commit skipped for phase {phase_no}.")
            if phase_no >= HINT_PHASE_NO:
                print("Hint/high-risk phase approved but left uncommitted for manual testing/review.")
                print("Stopping so Richard can manually test before committing.")
                break
        completed += 1
        phase_no += 1
    print("\n✅ Requested WES2 orchestration completed.")
    print("Logs:", log_dir)
    print("Current status:")
    print(run(["git", "status", "--short"], timeout=30).out)


if __name__ == "__main__":
    main()
